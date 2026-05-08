"""Geração de documentos .docx a partir de modelos com campos destacados em amarelo."""

import io
import re
from datetime import datetime

from copy import deepcopy

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from config import MODELOS_DIR

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_MESES = [
    "JANEIRO", "FEVEREIRO", "MARÇO", "ABRIL", "MAIO", "JUNHO",
    "JULHO", "AGOSTO", "SETEMBRO", "OUTUBRO", "NOVEMBRO", "DEZEMBRO",
]
_MESES_MIN = [m.lower() for m in _MESES]

# Expansão de abreviações comuns de logradouro
_ABBREV = {
    "ROD.": "rodovia", "R.": "rua", "AV.": "avenida",
    "ESTR.": "estrada", "TRAV.": "travessa", "AL.": "alameda",
    "PÇA.": "praça", "PC.": "praça", "VIA": "via",
}


def listar_modelos() -> list[dict]:
    return [
        {"nome": p.stem.upper(), "path": str(p)}
        for p in sorted(MODELOS_DIR.glob("*.docx"))
    ]


# ── Helpers de run ─────────────────────────────────────────────────────────────

def _get_highlight(run) -> str | None:
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return None
    h = rpr.find(qn("w:highlight"))
    if h is None:
        return None
    return h.get(qn("w:val"))


def _clear_highlight(run) -> None:
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return
    h = rpr.find(qn("w:highlight"))
    if h is not None:
        rpr.remove(h)


def _merge_yellow_runs(para) -> None:
    """Junta runs amarelos consecutivos sem cruzar quebras de linha."""
    runs = para.runs
    i = 0
    while i < len(runs) - 1:
        r1, r2 = runs[i], runs[i + 1]
        if (
            _get_highlight(r1) == "yellow"
            and _get_highlight(r2) == "yellow"
            and "\n" not in r1.text
            and "\n" not in r2.text
        ):
            r1.text += r2.text
            r2._r.getparent().remove(r2._r)
            runs = para.runs
        else:
            i += 1


# ── Helpers de formatação ──────────────────────────────────────────────────────

_PREPOSICOES = {"de", "da", "do", "das", "dos", "e", "a", "o", "em", "na", "no"}


def _title_case(text: str) -> str:
    """Converte para title case com preposições em minúsculas (exceto primeira palavra)."""
    words = text.lower().split()
    result = []
    for i, w in enumerate(words):
        result.append(w if (i > 0 and w in _PREPOSICOES) else w.capitalize())
    return " ".join(result)


def _expand_logradouro(log: str) -> str:
    """Expande abreviação do tipo de logradouro e aplica title case."""
    parts = log.split()
    if not parts:
        return log
    abbrev = parts[0].upper()
    if abbrev in _ABBREV:
        rest = " ".join(p.capitalize() for p in parts[1:])
        return f"{_ABBREV[abbrev]} {rest}"
    return _title_case(log)


# ── Mapa de substituições ──────────────────────────────────────────────────────

def _build_replacements(cliente: dict, tecnico: dict) -> dict:
    p = "end" if (cliente.get("obra_mesmo") or not cliente.get("obra_log")) else "obra"

    num    = cliente.get(f"{p}_num", "")
    comp   = cliente.get(f"{p}_comp", "")
    bairro = cliente.get(f"{p}_bairro", "")
    cidade = cliente.get(f"{p}_cidade", "")
    estado = cliente.get(f"{p}_estado", "")
    cep    = cliente.get(f"{p}_cep", "")
    cpf    = cliente.get("cpf", "")
    tipo   = cliente.get("tipo", "PF")
    now    = datetime.now()

    cpf_label = "CPF" if tipo == "PF" else "CNPJ"

    return {
        # ── Dados da área ──────────────────────────────────────────────
        "-LOGRADOURO / RUA-": cliente.get(f"{p}_log", ""),
        "-Nº-":               f"Nº {num}" if num else "",
        "-COMPLEMENTO-":      comp,
        "-BAIRRO-":           f"BAIRRO {bairro}" if bairro else "",
        "-CEP-":              f"CEP: {cep}" if cep else "",
        "-CIDADE-":           cidade,
        "-UF-":               estado,
        "-NOME DO CLIENTE-":  cliente.get("nome", ""),
        "-CPF/CNPJ-":         f"{cpf_label}: {cpf}",
        # ── Localização (será reconstruída por _fix_localizacao) ───────
        "-Bairro-":           bairro,
        "-*conseito numero*-": num,
        # ── Técnico ────────────────────────────────────────────────────
        "-NOME DO TÉCNICO-":  tecnico.get("nome", ""),
        "-NUMERO CFT-":       tecnico.get("cft", ""),
        # ── Capa (text boxes) ──────────────────────────────────────────
        "-MÊS-":  _MESES[now.month - 1],
        "-MES-":  _MESES[now.month - 1],
        "-ANO-":  str(now.year),
    }


# ── Substituição em parágrafos do body ────────────────────────────────────────

def _replace_in_paragraphs(paragraphs, replacements: dict) -> None:
    for para in paragraphs:
        _merge_yellow_runs(para)
        for run in para.runs:
            if _get_highlight(run) != "yellow":
                continue
            new_text = run.text
            for old, val in replacements.items():
                new_text = new_text.replace(old, val)
            if new_text != run.text:
                run.text = new_text
                _clear_highlight(run)


# ── Substituição em text boxes (capa) ─────────────────────────────────────────

def _replace_in_txbx(doc, replacements: dict) -> None:
    """Substituição por segmento (separado por w:br) dentro de cada text box.
    Processa todos os txbxContent — as substituições são idempotentes."""
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for p in txbx.findall(f".//{{{_W}}}p"):
            # Divide runs em segmentos separados por w:br
            segments: list[list] = []
            current: list = []
            for child in p:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag != "r":
                    continue
                br  = child.find(f"{{{_W}}}br")
                t   = child.find(f"{{{_W}}}t")
                has_text = t is not None and t.text
                if br is not None and not has_text:
                    segments.append(current)
                    current = []
                elif has_text:
                    current.append((child, t))
            if current:
                segments.append(current)

            # Substitui cada segmento independentemente
            for seg in segments:
                if not seg:
                    continue
                full = "".join(t.text or "" for _, t in seg)
                new  = full
                for old, val in replacements.items():
                    new = new.replace(old, val)
                if new != full:
                    seg[0][1].text = new
                    seg[0][1].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    for _, t in seg[1:]:
                        t.text = ""


# ── Limpeza de separadores duplos (ex: " -  - " quando campo vazio) ───────────

def _clean_address_separators(para, sep: str = " - ") -> None:
    """Remove runs separadores ' - ' adjacentes a runs vazios."""
    changed = True
    while changed:
        changed = False
        runs = para.runs
        for i, r in enumerate(runs):
            if r.text == "":
                # Tenta remover o separador anterior + este run vazio
                if i > 0 and runs[i - 1].text == sep:
                    r._r.getparent().remove(r._r)
                    runs[i - 1]._r.getparent().remove(runs[i - 1]._r)
                    changed = True
                    break
                # Ou o separador posterior + este run vazio
                elif i + 1 < len(runs) and runs[i + 1].text == sep:
                    runs[i + 1]._r.getparent().remove(runs[i + 1]._r)
                    r._r.getparent().remove(r._r)
                    changed = True
                    break


# ── Reconstroí o parágrafo de Localização com formatação adequada ─────────────

def _fix_localizacao(doc, cliente: dict, prefix: str) -> None:
    num    = cliente.get(f"{prefix}_num", "").strip()
    bairro = _title_case(cliente.get(f"{prefix}_bairro", ""))
    cidade = _title_case(cliente.get(f"{prefix}_cidade", ""))
    uf     = cliente.get(f"{prefix}_estado", "").upper()
    log_tc = _expand_logradouro(cliente.get(f"{prefix}_log", ""))

    if not num or num.upper() in ("S/N", "SN", "S/Nº", ""):
        num_text = "S/Nº"
    else:
        num_text = f"número {num}"

    part_before = f"A área está localizada no Bairro {bairro} em "
    part_city   = f"{cidade}/{uf}"
    part_after  = f", tendo acesso pela {log_tc}, {num_text}."

    for para in doc.paragraphs:
        if "localizada no Bairro" in para.text or (
            "tendo acesso pela" in para.text and "localizada" in para.text
        ):
            if not para.runs:
                break
            # Copia propriedades de formatação do primeiro run como base
            base_rpr = para.runs[0]._r.find(qn("w:rPr"))

            # Remove todos os runs existentes
            p_elem = para._p
            for r_elem in [r._r for r in para.runs]:
                p_elem.remove(r_elem)

            def _make_run(text: str, bold: bool) -> None:
                r = OxmlElement("w:r")
                rpr = deepcopy(base_rpr) if base_rpr is not None else OxmlElement("w:rPr")
                # Ajusta negrito
                b_tag = rpr.find(qn("w:b"))
                if bold:
                    if b_tag is None:
                        b_new = OxmlElement("w:b")
                        rpr.insert(0, b_new)
                else:
                    if b_tag is not None:
                        rpr.remove(b_tag)
                r.append(rpr)
                t = OxmlElement("w:t")
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                t.text = text
                r.append(t)
                p_elem.append(r)

            _make_run(part_before, bold=False)
            _make_run(part_city,   bold=False)
            _make_run(part_after,  bold=False)
            break


# ── Atualiza data no rodapé ────────────────────────────────────────────────────

def _update_date(doc) -> None:
    now = datetime.now()
    new_text = f"São José, {now.day} de {_MESES_MIN[now.month - 1]} de {now.year}."
    pattern = re.compile(r".+,\s+\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\.", re.IGNORECASE)
    for para in doc.paragraphs:
        if pattern.match(para.text.strip()):
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            break


# ── Corrige formatação do parágrafo Município ─────────────────────────────────

def _fix_municipio_bold(doc) -> None:
    """'Município:' em negrito; nome da cidade sem negrito."""
    for para in doc.paragraphs:
        if para.text.startswith("Município"):
            for run in para.runs:
                if "Município" in run.text:
                    run.bold = True
                elif run.text.strip():
                    run.bold = False
            break


# ── Remove negrito da cidade/UF nas text boxes (capa) ────────────────────────

def _fix_txbx_cidade_bold(doc, cidade: str, uf: str) -> None:
    """Remove negrito dos runs que contêm o nome da cidade nas text boxes."""
    busca = cidade.strip().upper()
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for r in txbx.iter(f"{{{_W}}}r"):
            t = r.find(f"{{{_W}}}t")
            if t is None or not t.text:
                continue
            if busca in t.text.upper():
                rpr = r.find(qn("w:rPr"))
                if rpr is not None:
                    b = rpr.find(qn("w:b"))
                    if b is not None:
                        rpr.remove(b)


# ── Remove realces amarelos restantes ─────────────────────────────────────────

def _clear_remaining_yellow(doc) -> None:
    """Remove highlight amarelo de todos os runs que ainda tenham (campos não preenchidos)."""
    def _clear_paras(paras):
        for para in paras:
            for run in para.runs:
                if _get_highlight(run) == "yellow":
                    _clear_highlight(run)

    _clear_paras(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _clear_paras(cell.paragraphs)

    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for r in txbx.iter(f"{{{_W}}}r"):
            rpr = r.find(qn("w:rPr"))
            if rpr is not None:
                h = rpr.find(qn("w:highlight"))
                if h is not None:
                    rpr.remove(h)


# ── Remove data bindings de SDTs (capa) ──────────────────────────────────────

def _remove_sdt_bindings(doc) -> None:
    """Remove w:dataBinding de content controls para que o texto estático prevaleça."""
    for txbx in doc.element.body.iter(qn("w:txbxContent")):
        for sdt in txbx.findall(f".//{{{_W}}}sdt"):
            sdt_pr = sdt.find(f"{{{_W}}}sdtPr")
            if sdt_pr is not None:
                data_binding = sdt_pr.find(f"{{{_W}}}dataBinding")
                if data_binding is not None:
                    sdt_pr.remove(data_binding)


# ── Entry point ───────────────────────────────────────────────────────────────

def gerar_documento_bytes(modelo_path: str, cliente: dict, tecnico: dict) -> bytes:
    doc = Document(modelo_path)
    replacements = _build_replacements(cliente, tecnico)
    p = "end" if (cliente.get("obra_mesmo") or not cliente.get("obra_log")) else "obra"

    # Substitui body paragraphs
    _replace_in_paragraphs(doc.paragraphs, replacements)

    # Substitui tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs, replacements)

    # Substitui text boxes (capa)
    _replace_in_txbx(doc, replacements)

    # Remove negrito da cidade/UF na capa
    _fix_txbx_cidade_bold(doc, cliente.get(f"{p}_cidade", ""), cliente.get(f"{p}_estado", ""))

    # Remove data bindings de SDTs para que o texto estático prevaleça
    _remove_sdt_bindings(doc)

    # Remove todos os realces amarelos restantes
    _clear_remaining_yellow(doc)

    # Limpeza de separadores vazios no parágrafo de endereço
    for para in doc.paragraphs:
        if para.text.startswith("Endere") and "-" in para.text:
            _clean_address_separators(para)
            break

    # Remove negrito da cidade no parágrafo Município
    _fix_municipio_bold(doc)

    # Reconstrói parágrafo de Localização com formatação adequada
    _fix_localizacao(doc, cliente, p)

    # Atualiza data
    _update_date(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
