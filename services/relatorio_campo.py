"""Geração de relatório de campo (.docx) a partir dos registros de ponto."""

import io
from calendar import monthrange
from collections import defaultdict
from datetime import date, timedelta

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

_MESES_PT = [
    "Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
    "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro",
]
_DIAS_PT = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]


def _parse_data(data_str: str) -> date | None:
    try:
        d, m, y = data_str.split("/")
        return date(int(y), int(m), int(d))
    except Exception:
        return None


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.upper().lstrip("#"))
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "DDDDDD") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color.upper().lstrip("#"))
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _cell_text(cell, text: str, *, bold=False, size=9,
               color: str | None = None, align="left") -> None:
    para = cell.paragraphs[0]
    para.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)


def gerar_relatorio_bytes(
    registros: list[dict],
    modo: str,      # "obra" | "funcionario"
    filtro: str,    # nome da obra ou funcionário
    ano_ini: int,
    mes_ini: int,
    ano_fim: int,
    mes_fim: int,
) -> bytes:
    data_ini = date(ano_ini, mes_ini, 1)
    data_fim = date(ano_fim, mes_fim, monthrange(ano_fim, mes_fim)[1])

    group_key = "usuario" if modo == "obra" else "obra"

    # Indexa: data → group_val → [registros]
    por_dia: dict[date, dict[str, list]] = defaultdict(lambda: defaultdict(list))
    for r in registros:
        match = (
            r.get("obra", "") == filtro
            if modo == "obra"
            else r.get("usuario", "") == filtro
        )
        if not match:
            continue
        d = _parse_data(r.get("data", ""))
        if d and data_ini <= d <= data_fim:
            por_dia[d][r.get(group_key, "—")].append(r)

    # ── Documento ─────────────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2)
        sec.bottom_margin = Cm(2)
        sec.left_margin   = Cm(2)
        sec.right_margin  = Cm(2)

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO DE CAMPO")
    run.bold = True
    run.font.size = Pt(14)

    sub_label = "Obra" if modo == "obra" else "Funcionário"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{sub_label}: {filtro}")
    run.font.size = Pt(11)

    if mes_ini == mes_fim and ano_ini == ano_fim:
        periodo = f"{_MESES_PT[mes_ini - 1]} de {ano_ini}"
    else:
        periodo = f"{_MESES_PT[mes_ini - 1]}/{ano_ini} a {_MESES_PT[mes_fim - 1]}/{ano_fim}"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Período: {periodo}")
    run.font.size = Pt(10)
    run.font.color.rgb = _rgb("888888")
    doc.add_paragraph()

    # ── Tabela ────────────────────────────────────────────────────────
    col3_hdr = "Funcionário" if modo == "obra" else "Obra"
    headers  = ["Data", "Dia", col3_hdr, "Modificador", "Check-in", "Check-out"]
    widths   = [Cm(2.8), Cm(1.4), Cm(4.5), Cm(3.0), Cm(2.4), Cm(2.4)]

    table = doc.add_table(rows=1, cols=6)

    # Cabeçalho da tabela
    hrow = table.rows[0]
    for i, (h, w) in enumerate(zip(headers, widths)):
        cell = hrow.cells[i]
        cell.width = w
        _set_cell_bg(cell, "1E3A5F")
        _set_cell_borders(cell, "1E3A5F")
        _cell_text(cell, h, bold=True, size=9, color="FFFFFF", align="center")

    # Linha para cada dia do período
    cur = data_ini
    while cur <= data_fim:
        is_wknd = cur.weekday() >= 5
        row_bg  = "F0F0F0" if is_wknd else "FFFFFF"
        data_s  = cur.strftime("%d/%m/%Y")
        dia_s   = _DIAS_PT[cur.weekday()]
        groups  = por_dia.get(cur, {})

        _MOD_STYLE = {
            "Diária":       ("EFF6FF", "BFDBFE", "1E40AF"),
            "Meia Diária":  ("FFFBEB", "FDE68A", "92400E"),
            "Levantamento": ("ECFDF5", "A7F3D0", "065F46"),
        }

        if not groups:
            row = table.add_row()
            for j, w in enumerate(widths):
                row.cells[j].width = w
                _set_cell_bg(row.cells[j], row_bg)
                _set_cell_borders(row.cells[j])
            muted = "AAAAAA" if is_wknd else "999999"
            _cell_text(row.cells[0], data_s, size=9, color=muted, align="center")
            _cell_text(row.cells[1], dia_s,  size=9, color="BBBBBB", align="center")
            _cell_text(row.cells[3], "—",    size=9, color="CCCCCC", align="center")
            _cell_text(row.cells[4], "—",    size=9, color="CCCCCC", align="center")
            _cell_text(row.cells[5], "—",    size=9, color="CCCCCC", align="center")
        else:
            first_row_for_date = True
            for grp, recs in sorted(groups.items()):
                recs_s = sorted(recs, key=lambda x: x.get("hora", ""))
                checkins  = [r for r in recs_s if r.get("tipo") == "checkin"]
                checkouts = [r for r in recs_s if r.get("tipo") == "checkout"]

                # Pair each checkin with the first checkout that comes after it
                pairs: list[tuple] = []
                used_co: set[int] = set()
                for ci in checkins:
                    ci_hora = ci.get("hora", "")
                    co_match = None
                    for idx, co in enumerate(checkouts):
                        if idx not in used_co and co.get("hora", "") >= ci_hora:
                            co_match = co
                            used_co.add(idx)
                            break
                    pairs.append((ci, co_match))

                # Unpaired checkouts (checkout without a preceding checkin)
                for idx, co in enumerate(checkouts):
                    if idx not in used_co:
                        pairs.append((None, co))

                # Fallback: no pairs at all (shouldn't happen, but guard)
                if not pairs:
                    pairs = [(None, None)]

                for ci, co in pairs:
                    mod = (ci or {}).get("modificador", "") or ""

                    row = table.add_row()
                    for j, w in enumerate(widths):
                        row.cells[j].width = w
                        _set_cell_bg(row.cells[j], row_bg)
                        _set_cell_borders(row.cells[j])

                    date_color = "AAAAAA" if is_wknd else "333333"
                    if first_row_for_date:
                        _cell_text(row.cells[0], data_s, bold=True, size=9, color=date_color, align="center")
                        _cell_text(row.cells[1], dia_s,  size=9, color="888888", align="center")
                        first_row_for_date = False

                    _cell_text(row.cells[2], grp, size=9)

                    if mod and mod in _MOD_STYLE:
                        bg, border, txt = _MOD_STYLE[mod]
                        _set_cell_bg(row.cells[3], bg)
                        _set_cell_borders(row.cells[3], border)
                        _cell_text(row.cells[3], mod, bold=True, size=8, color=txt, align="center")
                    else:
                        _cell_text(row.cells[3], "—", size=9, color="CCCCCC", align="center")

                    if ci:
                        _set_cell_bg(row.cells[4], "D1FAE5")
                        _set_cell_borders(row.cells[4], "A7F3D0")
                        _cell_text(row.cells[4], ci["hora"], bold=True, size=9, color="065F46", align="center")
                    else:
                        _cell_text(row.cells[4], "—", size=9, color="CCCCCC", align="center")

                    if co:
                        _set_cell_bg(row.cells[5], "DBEAFE")
                        _set_cell_borders(row.cells[5], "BFDBFE")
                        _cell_text(row.cells[5], co["hora"], bold=True, size=9, color="1E3A8A", align="center")
                    else:
                        _cell_text(row.cells[5], "—", size=9, color="CCCCCC", align="center")

        cur += timedelta(days=1)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
