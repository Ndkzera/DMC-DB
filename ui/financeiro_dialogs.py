"""Dialogs do módulo Financeiro / NFS-e."""

import base64
import json
from datetime import datetime, date
from pathlib import Path

from nicegui import ui

from services.nfse import (
    load_config, save_config, emit_nfse, save_nfse,
    cert_info, cancel_nfse, list_nfse, _DEPS_OK, _DEPS_MSG,
)


# ── Helpers de layout ──────────────────────────────────────────────────

def _section(label: str):
    ui.html(
        f'<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
        f'letter-spacing:.14em;text-transform:uppercase;'
        f'border-top:1px solid var(--dmc-b1);padding-top:14px;margin-bottom:10px">'
        f'{label}</div>'
    )


def _label(txt: str):
    ui.html(f'<label class="dmc-label">{txt}</label>')


def _row2(fn):
    with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:12px'):
        fn()


# ── Cadastro da Empresa ────────────────────────────────────────────────

def empresa_dialog(on_save=None) -> None:
    """Dialog completo de cadastro da empresa prestadora."""
    cfg = load_config()
    _st = dict(cfg)
    _ireg: dict = {}  # key → html input id (para leitura em batch via JS)

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(720px,97vw)!important;max-height:96vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        # ── Header ─────────────────────────────────────────────────────
        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(96,165,250,.08);border:1px solid rgba(96,165,250,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#60A5FA">business</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Cadastro da Empresa</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Dados do prestador de serviço — usados automaticamente na emissão de NFS-e</div>')

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0;display:flex;flex-direction:column;gap:0'):

            def _lbl(txt: str):
                ui.html(f'<label class="dmc-label">{txt}</label>')

            def _inp(label: str, key: str, placeholder: str = '', mono: bool = False):
                _lbl(label)
                inp_id = f'ei-{key}'
                _ireg[key] = inp_id
                val = str(_st.get(key, '') or '').replace('"', '&quot;')
                font_style = 'font-family:var(--dmc-mono);' if mono else ''
                ui.html(
                    f'<input id="{inp_id}" class="dmc-input" value="{val}" '
                    f'placeholder="{placeholder}" '
                    f'style="{font_style}margin-bottom:12px">'
                )

            async def _sync():
                """Lê valores atuais dos inputs nativos e sincroniza com _st."""
                if not _ireg:
                    return
                vals = await ui.run_javascript(
                    '(function(){var d=' + json.dumps(_ireg) + ',o={};'
                    'for(var k in d){var e=document.getElementById(d[k]);if(e)o[k]=e.value;}'
                    'return o;})()'
                )
                if isinstance(vals, dict):
                    _st.update(vals)

            # ── Seção 1: Identificação ──────────────────────────────────
            ui.html(
                '<div style="font:600 10px var(--dmc-mono);color:var(--dmc-muted2);letter-spacing:.14em;'
                'text-transform:uppercase;border-top:1px solid var(--dmc-b1);padding-top:14px;'
                'margin-bottom:14px;display:flex;align-items:center;gap:8px">'
                '<span class="material-icons" style="font-size:13px">badge</span>Identificação</div>'
            )

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    # CNPJ com botão de busca
                    _lbl('CNPJ')
                    cnpj_id = 'ei-cnpj'
                    _ireg['cnpj'] = cnpj_id
                    cnpj_val = str(_st.get('cnpj', '') or '').replace('"', '&quot;')
                    with ui.element('div').style('display:flex;gap:6px;margin-bottom:12px;align-items:center'):
                        ui.html(
                            f'<input id="{cnpj_id}" class="dmc-input" value="{cnpj_val}" '
                            f'placeholder="00.000.000/0000-00" '
                            f'style="font-family:var(--dmc-mono);flex:1;min-width:0">'
                        )
                        cnpj_status = ui.html('<span style="font:10px var(--dmc-mono);white-space:nowrap;flex-shrink:0"></span>')

                        async def _buscar_cnpj():
                            cnpj_now = await ui.run_javascript(
                                f'document.getElementById("{cnpj_id}")?.value||""'
                            )
                            _st['cnpj'] = cnpj_now or ''
                            raw = ''.join(c for c in (_st.get('cnpj') or '') if c.isdigit())
                            if len(raw) != 14:
                                ui.notify('CNPJ inválido (14 dígitos).', type='warning')
                                return
                            cnpj_status.set_content('<span style="color:#FBBF24">Consultando...</span>')
                            try:
                                import httpx
                                async with httpx.AsyncClient(timeout=10) as client:
                                    r = await client.get(
                                        f'https://www.receitaws.com.br/v1/cnpj/{raw}',
                                        headers={'Accept': 'application/json'},
                                    )
                                d = r.json()
                                if d.get('status') == 'ERROR':
                                    cnpj_status.set_content('<span style="color:#F87171">Não encontrado</span>')
                                    return
                                nome    = (d.get('nome') or '').strip().upper()
                                fantasia = (d.get('fantasia') or '').strip().upper()
                                tel     = (d.get('telefone') or '').split('/')[0].strip()
                                cep_r   = ''.join(c for c in (d.get('cep') or '') if c.isdigit())
                                await _sync()
                                _st.update({
                                    'razao_social':  nome,
                                    'nome_fantasia': fantasia,
                                    'telefone':      tel,
                                    'logradouro':    (d.get('logradouro') or '').upper(),
                                    'numero':        d.get('numero') or '',
                                    'complemento':   (d.get('complemento') or '').upper(),
                                    'bairro':        (d.get('bairro') or '').upper(),
                                    'cidade':        (d.get('municipio') or '').upper(),
                                    'uf':            (d.get('uf') or '').upper(),
                                    'cep_empresa':   cep_r[:5] + '-' + cep_r[5:] if len(cep_r) == 8 else '',
                                })
                                _render_fields()
                                cnpj_status.set_content('<span style="color:#4ADE80">✓ Preenchido</span>')
                            except Exception as ex:
                                cnpj_status.set_content(f'<span style="color:#F87171">Erro: {ex}</span>')

                        btn_cnpj = ui.element('button').classes('dmc-btn dmc-btn-secondary dmc-btn-sm').style('flex-shrink:0')
                        with btn_cnpj:
                            ui.html('<span class="material-icons" style="font-size:13px">search</span><span>Buscar</span>')
                        btn_cnpj.on('click', _buscar_cnpj)

                with ui.element('div'):
                    _inp('Inscrição Municipal', 'im', 'Nº municipal', mono=True)

            # área dos campos que podem ser preenchidos via busca CNPJ/CEP
            fields_area = ui.element('div')

            def _render_fields():
                fields_area.clear()
                with fields_area:
                    _inp('Razão Social', 'razao_social', 'Nome jurídico da empresa')

                    with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('Nome Fantasia', 'nome_fantasia', 'Nome comercial')
                        with ui.element('div'):
                            _inp('Inscrição Estadual', 'ie', 'Nº estadual ou ISENTO', mono=True)

                    with ui.element('div').style('margin-bottom:12px'):
                        _lbl('Regime Tributário')
                        reg_atual = _st.get('regime_trib', '')
                        with ui.element('div').style('display:flex;gap:8px;flex-wrap:wrap'):
                            for val, lbl in [
                                ('1', 'Simples Nacional'),
                                ('2', 'Lucro Presumido'),
                                ('3', 'Lucro Real'),
                                ('5', 'MEI'),
                            ]:
                                checked = 'checked' if reg_atual == val else ''
                                ui.html(
                                    f'<label style="display:inline-flex;align-items:center;gap:6px;cursor:pointer;'
                                    f'background:var(--dmc-bg3);border:1px solid var(--dmc-b2);border-radius:8px;'
                                    f'padding:7px 14px;font:12px var(--dmc-fm);color:var(--dmc-text);margin-bottom:12px">'
                                    f'<input type="radio" name="emp-regime" value="{val}" '
                                    f'style="accent-color:#60A5FA" {checked}> {lbl}</label>'
                                )

                    # ── Seção 2: Endereço ───────────────────────────────
                    ui.html(
                        '<div style="font:600 10px var(--dmc-mono);color:var(--dmc-muted2);letter-spacing:.14em;'
                        'text-transform:uppercase;border-top:1px solid var(--dmc-b1);padding-top:14px;'
                        'margin-bottom:14px;display:flex;align-items:center;gap:8px">'
                        '<span class="material-icons" style="font-size:13px">location_on</span>Endereço</div>'
                    )

                    with ui.element('div').style('display:flex;gap:8px;align-items:flex-end'):
                        with ui.element('div').style('flex:1;min-width:0'):
                            _inp('CEP', 'cep_empresa', '00000-000', mono=True)
                        cep_btn = ui.element('button').classes('dmc-btn dmc-btn-secondary dmc-btn-sm').style('margin-bottom:12px;flex-shrink:0')
                        with cep_btn:
                            ui.html('<span class="material-icons" style="font-size:13px">search</span><span>Buscar</span>')

                        async def _buscar_cep():
                            await _sync()
                            raw = ''.join(c for c in (_st.get('cep_empresa') or '') if c.isdigit())
                            if len(raw) != 8:
                                ui.notify('CEP inválido (8 dígitos).', type='warning')
                                return
                            try:
                                import httpx
                                async with httpx.AsyncClient(timeout=8) as client:
                                    r = await client.get(f'https://viacep.com.br/ws/{raw}/json/')
                                d = r.json()
                                if d.get('erro'):
                                    ui.notify('CEP não encontrado.', type='warning')
                                    return
                                await _sync()
                                _st.update({
                                    'logradouro': (d.get('logradouro') or '').upper(),
                                    'bairro':     (d.get('bairro') or '').upper(),
                                    'cidade':     (d.get('localidade') or '').upper(),
                                    'uf':         (d.get('uf') or '').upper(),
                                })
                                _render_fields()
                                ui.notify('Endereço preenchido.', type='positive')
                            except Exception as ex:
                                ui.notify(f'Erro ao buscar CEP: {ex}', type='negative')

                        cep_btn.on('click', _buscar_cep)

                    with ui.element('div').style('display:grid;grid-template-columns:2fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('Logradouro', 'logradouro', 'Rua / Avenida')
                        with ui.element('div'):
                            _inp('Número', 'numero', 'S/N')

                    with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('Complemento', 'complemento', 'Sala, bloco, andar...')
                        with ui.element('div'):
                            _inp('Bairro', 'bairro', '')

                    with ui.element('div').style('display:grid;grid-template-columns:2fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('Cidade', 'cidade', '')
                        with ui.element('div'):
                            _inp('UF', 'uf', 'SC')

                    # ── Seção 3: Contato ────────────────────────────────
                    ui.html(
                        '<div style="font:600 10px var(--dmc-mono);color:var(--dmc-muted2);letter-spacing:.14em;'
                        'text-transform:uppercase;border-top:1px solid var(--dmc-b1);padding-top:14px;'
                        'margin-bottom:14px;display:flex;align-items:center;gap:8px">'
                        '<span class="material-icons" style="font-size:13px">contact_mail</span>Contato</div>'
                    )

                    with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('E-mail da empresa', 'email_empresa', 'contato@empresa.com.br')
                        with ui.element('div'):
                            _inp('Telefone', 'telefone', '(00) 00000-0000', mono=True)

                    # ── Seção 4: Fiscal ─────────────────────────────────
                    ui.html(
                        '<div style="font:600 10px var(--dmc-mono);color:var(--dmc-muted2);letter-spacing:.14em;'
                        'text-transform:uppercase;border-top:1px solid var(--dmc-b1);padding-top:14px;'
                        'margin-bottom:14px;display:flex;align-items:center;gap:8px">'
                        '<span class="material-icons" style="font-size:13px">receipt_long</span>Fiscal / NFS-e</div>'
                    )

                    with ui.element('div').style('display:grid;grid-template-columns:2fr 1fr 1fr;gap:14px'):
                        with ui.element('div'):
                            _inp('Código IBGE do Município', 'cod_municipio', 'Ex: 4205407', mono=True)
                        with ui.element('div'):
                            _inp('Alíquota ISS (%)', 'aliquota_iss', '5.00', mono=True)
                        with ui.element('div'):
                            _inp('Série DPS', 'serie', '1', mono=True)

            _render_fields()

        # ── Rodapé ─────────────────────────────────────────────────────
        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            async def _salvar():
                await _sync()
                regime = await ui.run_javascript(
                    "document.querySelector('input[name=\"emp-regime\"]:checked')?.value||''"
                )
                _st['regime_trib'] = regime
                try:
                    cfg_atual = load_config()
                    cfg_atual.update(_st)
                    save_config(cfg_atual)
                    ui.notify('Dados da empresa salvos.', type='positive')
                    if on_save:
                        on_save()
                    dlg.close()
                except Exception as ex:
                    ui.notify(f'Erro ao salvar: {ex}', type='negative')

            ui.button('Salvar Empresa', icon='save', on_click=_salvar).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary')

    dlg.open()


# ── Certificado Digital ───────────────────────────────────────────────

def certificado_dialog(on_save=None) -> None:
    """Dialog profissional para gerenciamento do certificado digital A1."""
    cfg = load_config()
    _st = {'cert_path': cfg.get('cert_path', ''), 'cert_senha': cfg.get('cert_senha', '')}

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(520px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(167,139,250,.1);border:1px solid rgba(167,139,250,.3);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#A78BFA">verified_user</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Certificado Digital A1</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Upload, senha e validação do certificado .pfx / .p12</div>')

        with ui.element('div').style('padding:22px 24px;overflow-y:auto;flex:1;min-height:0'):

            status_card = ui.element('div').style('margin-bottom:18px')
            verify_card = ui.html('')

            def _render_status():
                status_card.clear()
                path = _st.get('cert_path', '')
                with status_card:
                    if path and Path(path).exists():
                        fname = Path(path).name
                        try:
                            kb = Path(path).stat().st_size / 1024
                            fsize = f'{kb:.1f} KB'
                        except Exception:
                            fsize = ''
                        ui.html(
                            f'<div style="background:rgba(167,139,250,.06);'
                            f'border:1px solid rgba(167,139,250,.25);border-radius:12px;'
                            f'padding:14px 18px;display:flex;align-items:center;gap:14px">'
                            f'<span class="material-icons" style="color:#A78BFA;font-size:28px;flex-shrink:0">lock</span>'
                            f'<div style="flex:1;min-width:0">'
                            f'<div style="font:600 13px var(--dmc-mono);color:var(--dmc-text);'
                            f'overflow:hidden;text-overflow:ellipsis;white-space:nowrap">{fname}</div>'
                            + (f'<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);margin-top:3px">{fsize}</div>' if fsize else '')
                            + f'</div>'
                            f'<span class="material-icons" style="color:#4ADE80;font-size:18px;flex-shrink:0">check_circle</span>'
                            f'</div>'
                        )
                    elif path:
                        fname = Path(path).name or path
                        ui.html(
                            f'<div style="background:rgba(248,113,113,.05);'
                            f'border:1px solid rgba(248,113,113,.22);border-radius:12px;'
                            f'padding:14px 18px;display:flex;align-items:center;gap:12px">'
                            f'<span class="material-icons" style="color:#F87171;font-size:22px;flex-shrink:0">error_outline</span>'
                            f'<div style="font:11px var(--dmc-mono);color:#F87171;overflow:hidden;'
                            f'text-overflow:ellipsis;white-space:nowrap">Arquivo não encontrado: {fname}</div>'
                            f'</div>'
                        )
                    else:
                        ui.html(
                            '<div style="background:var(--dmc-bg3);'
                            'border:1px solid var(--dmc-b1);border-radius:12px;'
                            'padding:14px 18px;display:flex;align-items:center;gap:12px">'
                            '<span class="material-icons" style="color:var(--dmc-muted2);font-size:22px;flex-shrink:0">lock_open</span>'
                            '<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2)">Nenhum certificado configurado</div>'
                            '</div>'
                        )

            _render_status()

            _section('Selecionar Certificado')

            _dz_id = f'cert-dz-{id(dlg)}'
            _fi_id = f'cert-fi-{id(dlg)}'

            ui.html(
                f'<div id="{_dz_id}"'
                f' style="border:2px dashed var(--dmc-b2);border-radius:12px;'
                f'background:var(--dmc-bg3);padding:28px 20px;'
                f'display:flex;flex-direction:column;align-items:center;gap:8px;'
                f'cursor:pointer;transition:border-color .15s;margin-bottom:14px">'
                f'<span class="material-icons" style="font-size:32px;color:rgba(167,139,250,.5)">upload_file</span>'
                f'<div style="font:600 12px var(--dmc-fd);color:var(--dmc-text)">Arraste ou clique para selecionar</div>'
                f'<div id="{_dz_id}-name" style="font:11px var(--dmc-mono);color:var(--dmc-muted2)">.pfx · .p12</div>'
                f'</div>'
                f'<input type="file" id="{_fi_id}" accept=".pfx,.p12" style="display:none">'
            )

            async def _on_cert_file(e):
                data = e.args if isinstance(e.args, dict) else (e.args[0] if e.args else {})
                name = data.get('name', 'cert.pfx')
                b64  = data.get('data', '')
                if not b64:
                    return
                try:
                    raw = base64.b64decode(b64)
                except Exception:
                    ui.notify('Erro ao decodificar o arquivo.', type='negative')
                    return
                certs_dir = Path('data') / 'certs'
                certs_dir.mkdir(parents=True, exist_ok=True)
                dest = certs_dir / name
                dest.write_bytes(raw)
                _st['cert_path'] = str(dest)
                _render_status()
                verify_card.set_content('')
                ui.notify(f'"{name}" carregado com sucesso.', type='positive')

            ui.on('cert_file_ready', _on_cert_file)

            async def _setup_dz():
                await ui.run_javascript(
                    f'(function(){{'
                    f'var dz=document.getElementById("{_dz_id}");'
                    f'var fi=document.getElementById("{_fi_id}");'
                    f'if(!dz||!fi)return;'
                    f'function hf(file){{'
                    f'if(!file)return;'
                    f'var nm=document.getElementById("{_dz_id}-name");'
                    f'if(nm)nm.textContent=file.name;'
                    f'var r=new FileReader();'
                    f'r.onload=function(ev){{'
                    f'var b64=ev.target.result.split(",")[1];'
                    f'emitEvent("cert_file_ready",{{name:file.name,data:b64,size:file.size}});'
                    f'}};'
                    f'r.readAsDataURL(file);'
                    f'}}'
                    f'dz.addEventListener("dragover",function(e){{e.preventDefault();dz.style.borderColor="#A78BFA";}});'
                    f'dz.addEventListener("dragleave",function(){{dz.style.borderColor="";}});'
                    f'dz.addEventListener("drop",function(e){{e.preventDefault();dz.style.borderColor="";hf(e.dataTransfer.files[0]);}});'
                    f'dz.addEventListener("click",function(){{fi.click();}});'
                    f'fi.addEventListener("change",function(){{if(fi.files[0])hf(fi.files[0]);}});'
                    f'}})();'
                )

            ui.timer(0.2, _setup_dz, once=True)

            _section('Senha do Certificado')
            _pw_id = f'cert-pw-{id(dlg)}'
            with ui.element('div').style('position:relative;margin-bottom:16px;width:100%'):
                ui.html(
                    f'<input type="password" id="{_pw_id}" class="dmc-input"'
                    f' value="{_st.get("cert_senha", "")}" placeholder="••••••••" autocomplete="off"'
                    f' style="padding-right:44px;height:40px;width:100%;box-sizing:border-box">'
                )
                with ui.element('button').style(
                    'position:absolute;right:0;top:0;height:40px;width:40px;'
                    'background:transparent;border:none;cursor:pointer;'
                    'display:flex;align-items:center;justify-content:center;padding:0;'
                ).props('type=button tabindex=-1') as _pw_tog:
                    _pw_icon = ui.html(
                        '<span class="material-icons" '
                        'style="font-size:18px;color:var(--dmc-muted)">visibility_off</span>'
                    )

                async def _pw_toggle():
                    state = await ui.run_javascript(
                        f'var i=document.getElementById("{_pw_id}");'
                        'if(!i)return "x";'
                        'if(i.type==="password"){i.type="text";return "text";}'
                        'i.type="password";return "password";'
                    )
                    _pw_icon.set_content(
                        '<span class="material-icons" style="font-size:18px;color:var(--dmc-muted)">'
                        + ('visibility' if state == 'text' else 'visibility_off') + '</span>'
                    )
                _pw_tog.on('click', _pw_toggle)

            async def _verify():
                path = _st.get('cert_path', '')
                senha = await ui.run_javascript(
                    f'document.getElementById("{_pw_id}")?.value||""'
                )
                _st['cert_senha'] = senha
                if not path or not Path(path).exists():
                    verify_card.set_content(
                        '<div style="background:rgba(248,113,113,.06);border:1px solid rgba(248,113,113,.22);'
                        'border-radius:8px;padding:10px 14px;font:11px var(--dmc-mono);color:#F87171;margin-bottom:14px">'
                        'Nenhum certificado carregado.</div>'
                    )
                    return
                try:
                    info = cert_info(path, senha)
                    verify_card.set_content(
                        f'<div style="background:rgba(74,222,128,.05);border:1px solid rgba(74,222,128,.2);'
                        f'border-radius:8px;padding:12px 16px;display:flex;gap:20px;flex-wrap:wrap;margin-bottom:14px">'
                        f'<div><div style="font:10px var(--dmc-mono);color:var(--dmc-muted2);'
                        f'text-transform:uppercase;letter-spacing:.08em;margin-bottom:3px">Titular</div>'
                        f'<div style="font:600 12px var(--dmc-mono);color:#A78BFA">{info["cn"]}</div></div>'
                        f'<div><div style="font:10px var(--dmc-mono);color:var(--dmc-muted2);'
                        f'text-transform:uppercase;letter-spacing:.08em;margin-bottom:3px">Válido até</div>'
                        f'<div style="font:600 12px var(--dmc-mono);color:#4ADE80">{info["validade"]}</div></div>'
                        f'</div>'
                    )
                except Exception as ex:
                    verify_card.set_content(
                        f'<div style="background:rgba(248,113,113,.06);border:1px solid rgba(248,113,113,.22);'
                        f'border-radius:8px;padding:10px 14px;font:11px var(--dmc-mono);color:#F87171;margin-bottom:14px">'
                        f'Erro: {ex}</div>'
                    )

            ui.button('Verificar certificado', icon='verified', on_click=_verify).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-secondary').style(
                'color:#A78BFA;border-color:rgba(167,139,250,.3)'
            )

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:space-between;align-items:center;gap:10px;flex-shrink:0'
        ):
            def _remove():
                _st['cert_path'] = ''
                _render_status()
                verify_card.set_content('')

            ui.button('Remover', icon='delete_outline', on_click=_remove).props(
                'flat no-caps'
            ).classes('dmc-btn').style('color:#F87171')

            with ui.element('div').style('display:flex;gap:10px'):
                ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

                async def _salvar():
                    senha = await ui.run_javascript(
                        f'document.getElementById("{_pw_id}")?.value||""'
                    )
                    _st['cert_senha'] = senha
                    try:
                        cfg2 = load_config()
                        cfg2['cert_path']  = _st['cert_path']
                        cfg2['cert_senha'] = _st['cert_senha']
                        save_config(cfg2)
                        ui.notify('Certificado salvo.', type='positive')
                        if on_save:
                            on_save()
                        dlg.close()
                    except Exception as ex:
                        ui.notify(f'Erro ao salvar: {ex}', type='negative')

                ui.button('Salvar', icon='save', on_click=_salvar).props(
                    'unelevated no-caps'
                ).classes('dmc-btn dmc-btn-primary')

    dlg.open()


# ── Configuração ───────────────────────────────────────────────────────

def config_nfse_dialog(on_save=None) -> None:
    cfg = load_config()
    _st = dict(cfg)
    _ireg: dict = {}

    async def _sync():
        if not _ireg:
            return
        vals = await ui.run_javascript(
            '(function(){var d=' + json.dumps(_ireg) + ',o={};'
            'for(var k in d){var e=document.getElementById(d[k]);if(e)o[k]=e.value;}return o;})()'
        )
        if isinstance(vals, dict):
            _st.update(vals)

    with ui.dialog() as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(680px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(74,222,128,.08);border:1px solid rgba(74,222,128,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#4ADE80">settings</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Configuração NFS-e</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Dados do prestador e certificado digital</div>')

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0'):

            ui.html('<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);letter-spacing:.14em;text-transform:uppercase;margin-bottom:10px">Prestador</div>')

            def _inp(label, key, placeholder='', mono=False):
                _label(label)
                font = 'var(--dmc-mono)' if mono else 'var(--dmc-fm)'
                _id = f'cfgn-{key}'
                _ireg[key] = _id
                val = str(_st.get(key, '') or '')
                ui.html(
                    f'<input id="{_id}" class="dmc-input" value="{val}"'
                    f' placeholder="{placeholder}"'
                    f' style="font:{font};margin-bottom:12px">'
                )

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    _inp('CNPJ', 'cnpj', '00.000.000/0000-00', mono=True)
                with ui.element('div'):
                    _inp('Inscrição Municipal', 'im', 'Nº municipal', mono=True)

            _inp('Razão Social', 'razao_social', 'Nome da empresa')

            with ui.element('div').style('display:grid;grid-template-columns:2fr 1fr;gap:14px'):
                with ui.element('div'):
                    _inp('Código IBGE do Município', 'cod_municipio', 'Ex: 4205407', mono=True)
                with ui.element('div'):
                    _inp('Série DPS', 'serie', '1', mono=True)

            _inp('Alíquota ISS (%)', 'aliquota_iss', '5.00', mono=True)

            _section('Ambiente')
            ui.html(
                '<div style="display:flex;gap:10px;margin-bottom:14px">'
                '<label style="display:inline-flex;align-items:center;gap:8px;cursor:pointer;'
                'background:var(--dmc-bg3);border:1px solid var(--dmc-b2);border-radius:8px;'
                'padding:8px 16px;font:13px var(--dmc-fm);color:var(--dmc-text)">'
                '<input type="radio" name="cfg-amb" value="homologacao" style="accent-color:#FBBF24"'
                + (' checked' if _st.get('ambiente') == 'homologacao' else '') +
                '> Homologação (testes)</label>'
                '<label style="display:inline-flex;align-items:center;gap:8px;cursor:pointer;'
                'background:var(--dmc-bg3);border:1px solid var(--dmc-b2);border-radius:8px;'
                'padding:8px 16px;font:13px var(--dmc-fm);color:var(--dmc-text)">'
                '<input type="radio" name="cfg-amb" value="producao" style="accent-color:#4ADE80"'
                + (' checked' if _st.get('ambiente') == 'producao' else '') +
                '> Produção</label>'
                '</div>'
            )

            _section('Certificado Digital')

            cert_card = ui.element('div').style('margin-bottom:4px')

            def _render_cert_card():
                cert_card.clear()
                path = _st.get('cert_path', '')
                with cert_card:
                    if path and Path(path).exists():
                        fname = Path(path).name
                        try:
                            kb = Path(path).stat().st_size / 1024
                            fsize = f'{kb:.1f} KB'
                        except Exception:
                            fsize = ''
                        ui.html(
                            f'<div style="background:rgba(167,139,250,.06);'
                            f'border:1px solid rgba(167,139,250,.25);border-radius:10px;'
                            f'padding:12px 16px;display:flex;align-items:center;gap:12px">'
                            f'<span class="material-icons" style="color:#A78BFA;font-size:22px;flex-shrink:0">lock</span>'
                            f'<div style="flex:1;min-width:0">'
                            f'<div style="font:600 12px var(--dmc-mono);color:var(--dmc-text);'
                            f'overflow:hidden;text-overflow:ellipsis;white-space:nowrap">{fname}</div>'
                            + (f'<div style="font:10px var(--dmc-mono);color:var(--dmc-muted2);margin-top:2px">{fsize}</div>' if fsize else '')
                            + f'</div>'
                            f'<span class="material-icons" style="color:#4ADE80;font-size:16px;flex-shrink:0">check_circle</span>'
                            f'</div>'
                        )
                    elif path:
                        fname = Path(path).name or path
                        ui.html(
                            f'<div style="background:rgba(248,113,113,.05);border:1px solid rgba(248,113,113,.22);'
                            f'border-radius:10px;padding:12px 16px;display:flex;align-items:center;gap:10px">'
                            f'<span class="material-icons" style="color:#F87171;font-size:18px;flex-shrink:0">error</span>'
                            f'<div style="font:11px var(--dmc-mono);color:#F87171;overflow:hidden;'
                            f'text-overflow:ellipsis;white-space:nowrap">Arquivo não encontrado: {fname}</div>'
                            f'</div>'
                        )
                    else:
                        ui.html(
                            '<div style="background:var(--dmc-bg3);border:1px solid var(--dmc-b1);'
                            'border-radius:10px;padding:12px 16px;display:flex;align-items:center;gap:10px">'
                            '<span class="material-icons" style="color:var(--dmc-muted2);font-size:18px;flex-shrink:0">lock_open</span>'
                            '<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2)">'
                            'Nenhum certificado · use o botão <b style="color:rgba(167,139,250,.9)">Certificado</b> na barra superior'
                            '</div></div>'
                        )

            _render_cert_card()

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            async def _salvar():
                await _sync()
                amb = await ui.run_javascript(
                    "document.querySelector('input[name=\"cfg-amb\"]:checked')?.value || 'homologacao'"
                )
                _st['ambiente'] = amb
                try:
                    save_config(_st)
                    ui.notify('Configuração salva.', type='positive')
                    if on_save:
                        on_save()
                    dlg.close()
                except Exception as ex:
                    ui.notify(f'Erro ao salvar: {ex}', type='negative')

            ui.button('Salvar', icon='save', on_click=_salvar).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary')

    dlg.open()


# ── Emissão NFS-e ──────────────────────────────────────────────────────

def emitir_nfse_dialog(on_success=None) -> None:
    cfg = load_config()

    if not cfg.get('cnpj'):
        ui.notify('Configure o CNPJ do prestador antes de emitir.', type='warning')
        config_nfse_dialog(on_save=lambda: emitir_nfse_dialog(on_success))
        return

    _st: dict = {
        'toma_tipo':   'CNPJ',
        'toma_doc':    '',
        'toma_nome':   '',
        'toma_cep':    '',
        'toma_end':    '',
        'toma_num':    '',
        'toma_bairro': '',
        'toma_cidade': '',
        'toma_uf':     '',
        'toma_cod_mun': '',
        'descricao':   '',
        'cod_tributacao': '010101',
        'cod_nbs':     '',
        'valor':       '',
        'iss_retido':  False,
    }
    _ireg: dict = {}

    async def _sync():
        if not _ireg:
            return
        vals = await ui.run_javascript(
            '(function(){var d=' + json.dumps(_ireg) + ',o={};'
            'for(var k in d){'
            'var e=document.getElementById(d[k]);'
            'if(e)o[k]=(e.tagName==="TEXTAREA"?e.value:e.value);}return o;})()'
        )
        if isinstance(vals, dict):
            _st.update(vals)

    ui.add_body_html(
        '<script>if(!window._nfseISSCalc){window._nfseISSCalc=function(v,aliq,pid){'
        'var p=document.getElementById(pid);if(!p)return;'
        'v=parseFloat(v.replace(",","."));'
        'if(!isNaN(v)&&v>0){'
        'var i=Math.round(v*aliq/100*100)/100;'
        'p.innerHTML="ISS ("+aliq.toFixed(2)+"%)&nbsp;<span style=\'color:#FBBF24\'>R$\xa0"+i.toFixed(2)+"</span>"'
        '+"&nbsp;&middot;&nbsp;L&iacute;q&nbsp;<span style=\'color:#4ADE80\'>R$\xa0"+(v-i).toFixed(2)+"</span>";'
        '}else{p.innerHTML="";}};}</script>'
    )

    with ui.dialog() as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(760px,97vw)!important;height:94vh;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(74,222,128,.08);border:1px solid rgba(74,222,128,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#4ADE80">receipt_long</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Emitir NFS-e</div>')
                amb = cfg.get('ambiente', 'homologacao')
                amb_col = '#FBBF24' if amb == 'homologacao' else '#4ADE80'
                amb_lbl = 'HOMOLOGAÇÃO' if amb == 'homologacao' else 'PRODUÇÃO'
                ui.html(
                    f'<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">'
                    f'Prestador: {cfg.get("razao_social") or cfg.get("cnpj")} · '
                    f'<span style="color:{amb_col};font-weight:600">{amb_lbl}</span></div>'
                )

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0'):

            # ── Card do Prestador (readonly) ────────────────────────────
            cnpj_d = cfg.get('cnpj') or '—'
            im_d   = cfg.get('im') or '—'
            rs_d   = cfg.get('razao_social') or '—'
            nf_d   = cfg.get('nome_fantasia') or ''
            ui.html(
                '<div style="background:rgba(74,222,128,.04);border:1px solid rgba(74,222,128,.15);'
                'border-radius:10px;padding:12px 16px;margin-bottom:16px">'
                '<div style="font:600 9px var(--dmc-mono);color:var(--dmc-muted2);text-transform:uppercase;'
                'letter-spacing:.12em;margin-bottom:10px;display:flex;align-items:center;gap:6px">'
                '<span class="material-icons" style="font-size:12px;color:#4ADE80">business</span>'
                'Prestador (sua empresa)</div>'
                '<div style="display:grid;grid-template-columns:auto auto 1fr;gap:8px 28px;align-items:start">'
                # CNPJ
                '<div>'
                '<div style="font:9px var(--dmc-mono);color:var(--dmc-muted2);text-transform:uppercase;'
                'letter-spacing:.06em;margin-bottom:3px">CNPJ</div>'
                f'<div style="font:600 13px var(--dmc-mono);color:var(--dmc-text)">{cnpj_d}</div>'
                '</div>'
                # IM
                '<div>'
                '<div style="font:9px var(--dmc-mono);color:var(--dmc-muted2);text-transform:uppercase;'
                'letter-spacing:.06em;margin-bottom:3px">Insc. Municipal</div>'
                f'<div style="font:600 13px var(--dmc-mono);color:var(--dmc-text)">{im_d}</div>'
                '</div>'
                # Razão Social
                '<div>'
                '<div style="font:9px var(--dmc-mono);color:var(--dmc-muted2);text-transform:uppercase;'
                'letter-spacing:.06em;margin-bottom:3px">Razão Social</div>'
                f'<div style="font:500 13px var(--dmc-fm);color:var(--dmc-text)">{rs_d}'
                + (f'<span style="font:11px var(--dmc-fm);color:var(--dmc-muted2);margin-left:8px">'
                   f'({nf_d})</span>' if nf_d else '')
                + '</div></div></div></div>'
            )

            # ── Tomador ────────────────────────────────────────────────
            ui.html(
                '<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
                'letter-spacing:.14em;text-transform:uppercase;margin-bottom:10px">Tomador</div>'
            )

            ui.html(
                '<div style="display:flex;gap:10px;margin-bottom:12px">'
                '<label style="display:inline-flex;align-items:center;gap:6px;cursor:pointer;'
                'font:13px var(--dmc-fm);color:var(--dmc-text)">'
                '<input type="radio" name="toma-tipo" value="CNPJ" checked style="accent-color:#4ADE80"> '
                'CNPJ</label>'
                '<label style="display:inline-flex;align-items:center;gap:6px;cursor:pointer;'
                'font:13px var(--dmc-fm);color:var(--dmc-text)">'
                '<input type="radio" name="toma-tipo" value="CPF" style="accent-color:#4ADE80"> '
                'CPF (Pessoa Física)</label>'
                '</div>'
            )

            def _f(label, key, placeholder='', mono=False, grid_span=False):
                font = 'var(--dmc-mono)' if mono else 'var(--dmc-fm)'
                _label(label)
                _id = f'nfse-f-{key}'
                _ireg[key] = _id
                ui.html(
                    f'<input id="{_id}" class="dmc-input" placeholder="{placeholder}"'
                    f' style="font:{font};margin-bottom:12px">'
                )

            with ui.element('div').style('display:grid;grid-template-columns:1fr 2fr;gap:14px'):
                with ui.element('div'):
                    _f('CPF / CNPJ', 'toma_doc', '', mono=True)
                with ui.element('div'):
                    _f('Razão Social / Nome', 'toma_nome', 'Nome do tomador')

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr 1fr 1fr;gap:14px'):
                with ui.element('div').style('grid-column:span 2'):
                    _f('Logradouro', 'toma_end', 'Rua/Avenida')
                with ui.element('div'):
                    _f('Número', 'toma_num', 'S/N')
                with ui.element('div'):
                    _f('CEP', 'toma_cep', '00000-000', mono=True)

            with ui.element('div').style('display:grid;grid-template-columns:2fr 2fr 1fr;gap:14px'):
                with ui.element('div'):
                    _f('Bairro', 'toma_bairro', '')
                with ui.element('div'):
                    _f('Cidade', 'toma_cidade', '')
                with ui.element('div'):
                    _f('UF', 'toma_uf', 'SC')

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    _f('Código IBGE do Tomador', 'toma_cod_mun', cfg.get('cod_municipio', ''), mono=True)

            # ── Serviço ────────────────────────────────────────────────
            _section('Serviço')

            _label('Descrição do serviço prestado')
            _ireg['descricao'] = 'nfse-f-descricao'
            ui.html(
                '<textarea id="nfse-f-descricao" class="dmc-input" rows="3"'
                ' placeholder="Descreva o serviço prestado..."'
                ' style="padding:8px 12px;resize:vertical;height:auto;margin-bottom:12px"></textarea>'
            )

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    _label('Código Tributação Nacional (cTribNac)')
                    _ireg['cod_tributacao'] = 'nfse-f-cod_tributacao'
                    ui.html(
                        '<input id="nfse-f-cod_tributacao" class="dmc-input" value="010101"'
                        ' placeholder="010101" style="font:var(--dmc-mono);margin-bottom:12px">'
                    )
                with ui.element('div'):
                    _label('Código NBS (cNBS — obrigatório 2026+)')
                    _ireg['cod_nbs'] = 'nfse-f-cod_nbs'
                    ui.html(
                        '<input id="nfse-f-cod_nbs" class="dmc-input"'
                        ' placeholder="Ex: 1.0301.00.00" style="font:var(--dmc-mono);margin-bottom:12px">'
                    )

            # Código NBS info
            ui.html(
                '<div style="font:10px var(--dmc-mono);color:var(--dmc-muted2);'
                'background:rgba(96,165,250,.06);border:1px solid rgba(96,165,250,.15);'
                'border-radius:6px;padding:6px 10px;margin-bottom:14px">'
                'Consulte o código NBS em: '
                '<span style="color:#60A5FA">nfse.gov.br → Nomenclatura Brasileira de Serviços</span>'
                '</div>'
            )

            # ── Valores ────────────────────────────────────────────────
            _section('Valores')

            aliq = cfg.get('aliquota_iss', '5.00')
            aliq_f = float(aliq or '5.00')
            _val_id  = f'nfse-val-{id(dlg)}'
            _iss_id  = f'nfse-iss-{id(dlg)}'
            _ireg['valor'] = _val_id

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    _label('Valor do serviço (R$)')
                    _oninput = f"_nfseISSCalc(this.value,{aliq_f},'{_iss_id}')"
                    ui.html(
                        f'<input id="{_val_id}" class="dmc-input" placeholder="0,00"'
                        f' oninput="{_oninput}"'
                        f' style="font:var(--dmc-mono);margin-bottom:4px">'
                    )
                    ui.html(
                        f'<div id="{_iss_id}" style="font:11px var(--dmc-mono);'
                        f'color:var(--dmc-muted2);min-height:18px;margin-bottom:8px"></div>'
                    )

                with ui.element('div'):
                    _label(f'Alíquota ISS configurada')
                    ui.html(
                        f'<div style="background:var(--dmc-bg3);border:1px solid var(--dmc-b2);'
                        f'border-radius:8px;padding:0 12px;height:40px;'
                        f'display:flex;align-items:center;margin-bottom:8px">'
                        f'<span style="font:600 15px var(--dmc-mono);color:#FBBF24">{aliq}%</span>'
                        f'<span style="font:11px var(--dmc-fm);color:var(--dmc-muted2);margin-left:8px">'
                        f'(configurável em ⚙ Configurar)</span>'
                        f'</div>'
                    )

            ui.html(
                '<label style="display:inline-flex;align-items:center;gap:8px;cursor:pointer;'
                'font:13px var(--dmc-fm);color:var(--dmc-text);margin-bottom:14px">'
                '<input type="checkbox" id="nfse-iss-retido" style="accent-color:#4ADE80;'
                'width:16px;height:16px"> ISS Retido pelo Tomador</label>'
            )

            status_box = ui.element('div')

        # ── Rodapé ─────────────────────────────────────────────────────
        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            btn_emitir = ui.button('Emitir NFS-e', icon='send').props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary')

            async def _emitir():
                if not _DEPS_OK:
                    ui.notify(_DEPS_MSG, type='negative', multi_line=True)
                    return

                await _sync()

                # Coleta dados do JS
                tipo = await ui.run_javascript(
                    "document.querySelector('input[name=\"toma-tipo\"]:checked')?.value || 'CNPJ'"
                )
                iss_ret = await ui.run_javascript(
                    "document.getElementById('nfse-iss-retido')?.checked || false"
                )
                _st['toma_tipo']  = tipo
                _st['iss_retido'] = bool(iss_ret)

                # Validações básicas
                erros = []
                if not _st.get('toma_doc'):
                    erros.append('CPF/CNPJ do tomador é obrigatório.')
                if not _st.get('toma_nome'):
                    erros.append('Nome do tomador é obrigatório.')
                if not _st.get('descricao'):
                    erros.append('Descrição do serviço é obrigatória.')
                if not _st.get('valor'):
                    erros.append('Valor do serviço é obrigatório.')
                try:
                    float(str(_st.get('valor', '')).replace(',', '.'))
                except ValueError:
                    erros.append('Valor inválido.')

                if erros:
                    status_box.clear()
                    with status_box:
                        ui.html(
                            '<div style="margin:0 0 8px;padding:10px 14px;'
                            'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.3);'
                            'border-radius:8px;font:12px var(--dmc-fm);color:#F87171">'
                            + '<br>'.join(erros) + '</div>'
                        )
                    return

                _st['valor'] = float(str(_st['valor']).replace(',', '.'))
                cfg_atual = load_config()
                cfg_atual['proximo_num'] = cfg_atual.get('proximo_num', 1)

                btn_emitir.disable()
                status_box.clear()
                with status_box:
                    ui.html(
                        '<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);'
                        'padding:8px 0">⏳ Enviando para a API NFS-e Nacional...</div>'
                    )

                try:
                    result = emit_nfse(_st, cfg_atual)
                    save_nfse(result)

                    if result['sucesso']:
                        cfg_atual['proximo_num'] = cfg_atual.get('proximo_num', 1) + 1
                        save_config(cfg_atual)
                        ui.notify(
                            f'NFS-e emitida com sucesso! Nº {result["numero"]}',
                            type='positive',
                        )
                        if on_success:
                            on_success()
                        dlg.close()
                    else:
                        erros_api = result.get('erros', ['Erro desconhecido.'])
                        status_box.clear()
                        with status_box:
                            ui.html(
                                '<div style="margin:8px 0;padding:10px 14px;'
                                'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.3);'
                                'border-radius:8px;font:12px var(--dmc-mono);color:#F87171">'
                                '<b>Erros da API:</b><br>'
                                + '<br>'.join(erros_api) + '</div>'
                            )
                        save_nfse(result)
                        if on_success:
                            on_success()
                except Exception as ex:
                    status_box.clear()
                    with status_box:
                        ui.html(
                            f'<div style="margin:8px 0;padding:10px 14px;'
                            f'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.3);'
                            f'border-radius:8px;font:12px var(--dmc-mono);color:#F87171">'
                            f'<b>Exceção:</b> {ex}</div>'
                        )
                finally:
                    btn_emitir.enable()

            btn_emitir.on('click', _emitir)

    dlg.open()


# ── Visualização NFS-e ─────────────────────────────────────────────────

def ver_nfse_dialog(entry: dict) -> None:
    sucesso = entry.get('sucesso', False)
    num     = entry.get('numero', '—')
    toma    = entry.get('tomador', '—')
    val     = entry.get('valor', 0)
    amb     = entry.get('ambiente', 'homologacao')
    emitido = (entry.get('emitido_em', '') or '')[:16].replace('T', ' ')
    chave   = entry.get('chave_acesso', '—')
    erros   = entry.get('erros', [])
    dados   = entry.get('dados', {})

    amb_col = '#FBBF24' if amb == 'homologacao' else '#4ADE80'
    amb_lbl = 'HOMOLOGAÇÃO' if amb == 'homologacao' else 'PRODUÇÃO'
    val_fmt = f'R$ {val:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')

    with ui.dialog() as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(680px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        status_color = '#4ADE80' if sucesso else '#F87171'
        status_icon  = 'check_circle' if sucesso else 'error'

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                f'<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                f'background:{status_color}14;border:1px solid {status_color}33;'
                f'display:flex;align-items:center;justify-content:center;">'
                f'<span class="material-icons" style="font-size:20px;color:{status_color}">{status_icon}</span></div>'
            )
            with ui.element('div'):
                ui.html(f'<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">NFS-e #{num}</div>')
                ui.html(
                    f'<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">'
                    f'<span style="color:{amb_col};font-weight:600">{amb_lbl}</span>'
                    f' · Emitida em {emitido}</div>'
                )

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0'):

            def _row(lbl, val, color='var(--dmc-text)'):
                ui.html(
                    f'<div style="display:flex;justify-content:space-between;'
                    f'padding:8px 0;border-bottom:1px solid var(--dmc-b1)">'
                    f'<span style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
                    f'text-transform:uppercase;letter-spacing:.06em">{lbl}</span>'
                    f'<span style="font:13px var(--dmc-mono);color:{color}">{val}</span>'
                    f'</div>'
                )

            _row('Número', f'#{num}', '#4ADE80')
            _row('Tomador', toma)
            _row('Valor', val_fmt, '#FBBF24')
            _row('Ambiente', amb_lbl, amb_col)
            _row('Emitido em', emitido)
            if chave and chave != '—':
                _row('Chave de acesso', chave[:20] + '...' if len(chave) > 20 else chave)
            if dados.get('descricao'):
                _row('Serviço', dados['descricao'][:80] + ('...' if len(dados.get('descricao', '')) > 80 else ''))
            if dados.get('cod_nbs'):
                _row('Código NBS', dados['cod_nbs'])

            if erros:
                ui.html(
                    '<div style="margin-top:14px;padding:10px 14px;'
                    'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.3);'
                    'border-radius:8px">'
                    '<div style="font:600 10px var(--dmc-mono);color:#F87171;'
                    'letter-spacing:.1em;text-transform:uppercase;margin-bottom:6px">Erros</div>'
                    + ''.join(f'<div style="font:12px var(--dmc-mono);color:#FCA5A5">{e}</div>' for e in erros)
                    + '</div>'
                )

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            if entry.get('nfse_xml'):
                async def _dl_xml():
                    xml_b64 = base64.b64encode(entry['nfse_xml'].encode()).decode()
                    fname   = f'NFSe_{num}.xml'
                    await ui.run_javascript(f'''
                        const a = document.createElement('a');
                        a.href = 'data:application/xml;base64,{xml_b64}';
                        a.download = '{fname}';
                        document.body.appendChild(a); a.click();
                        document.body.removeChild(a);
                    ''')
                ui.button('Download XML', icon='download', on_click=_dl_xml).props(
                    'unelevated no-caps'
                ).classes('dmc-btn dmc-btn-secondary')

            ui.button('Fechar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

    dlg.open()


# ── Relatório Financeiro ───────────────────────────────────────────────

def _fmt_brl(v: float) -> str:
    return f'R$ {v:,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')


def _filter_entries(entries: list, periodo: str, de: str, ate: str,
                    inc_emitida: bool, inc_homo: bool, inc_erro: bool) -> list:
    hoje = date.today()

    if periodo == 'mes_atual':
        ini = hoje.replace(day=1).isoformat()
        fim = hoje.isoformat()
    elif periodo == 'mes_anterior':
        if hoje.month == 1:
            ini = date(hoje.year - 1, 12, 1).isoformat()
            fim = date(hoje.year - 1, 12, 31).isoformat()
        else:
            import calendar
            last = calendar.monthrange(hoje.year, hoje.month - 1)[1]
            ini = date(hoje.year, hoje.month - 1, 1).isoformat()
            fim = date(hoje.year, hoje.month - 1, last).isoformat()
    elif periodo == '3meses':
        from dateutil.relativedelta import relativedelta  # type: ignore
        ini = (hoje - relativedelta(months=3)).isoformat()
        fim = hoje.isoformat()
    elif periodo == 'ano_atual':
        ini = hoje.replace(month=1, day=1).isoformat()
        fim = hoje.isoformat()
    else:
        ini = de or '2000-01-01'
        fim = ate or hoje.isoformat()

    out = []
    for e in entries:
        dt = (e.get('emitido_em', '') or '')[:10]
        if dt and not (ini <= dt <= fim):
            continue
        sucesso = e.get('sucesso', False)
        amb = e.get('ambiente', 'homologacao')
        if sucesso and amb == 'producao' and not inc_emitida:
            continue
        if sucesso and amb == 'homologacao' and not inc_homo:
            continue
        if not sucesso and not inc_erro:
            continue
        out.append(e)
    return out


def _build_report_html(
    entries: list, cfg: dict, periodo_label: str,
    contas_pagar: list | None = None,
    contas_receber: list | None = None,
) -> str:
    empresa  = cfg.get('razao_social') or 'Empresa'
    cnpj_raw = cfg.get('cnpj') or ''
    cnpj_fmt = (
        f'{cnpj_raw[:2]}.{cnpj_raw[2:5]}.{cnpj_raw[5:8]}/{cnpj_raw[8:12]}-{cnpj_raw[12:]}'
        if len(cnpj_raw) == 14 else cnpj_raw
    )
    now_str  = datetime.now().strftime('%d/%m/%Y %H:%M')
    aliq     = float(cfg.get('aliquota_iss', '5.00') or '5.00')

    # ── NFS-e ──
    total     = sum(float(e.get('valor', 0) or 0) for e in entries)
    total_iss = round(total * aliq / 100, 2)
    total_liq = round(total - total_iss, 2)
    n_emit    = sum(1 for e in entries if e.get('sucesso') and e.get('ambiente') == 'producao')
    n_homo    = sum(1 for e in entries if e.get('sucesso') and e.get('ambiente') == 'homologacao')
    n_erro    = sum(1 for e in entries if not e.get('sucesso'))

    rows_nfse = ''
    for e in entries:
        sucesso = e.get('sucesso', False)
        amb     = e.get('ambiente', 'homologacao')
        s_lbl, s_cls = ('ERRO','err') if not sucesso else (('EMITIDA','ok') if amb=='producao' else ('HOMOL.','hom'))
        val     = float(e.get('valor', 0) or 0)
        iss_v   = round(val * aliq / 100, 2)
        emitido = (e.get('emitido_em','') or '')[:10]
        desc    = (e.get('descricao','') or '')[:55] + ('…' if len(e.get('descricao','') or '') > 55 else '')
        rows_nfse += (
            f'<tr>'
            f'<td class="tc num-col">#{e.get("numero","—")}</td>'
            f'<td>{e.get("tomador","—")}<div class="desc">{desc}</div></td>'
            f'<td class="tr mono">{_fmt_brl(val)}</td>'
            f'<td class="tr mono muted">{_fmt_brl(iss_v)}</td>'
            f'<td class="tc"><span class="badge {s_cls}">{s_lbl}</span></td>'
            f'<td class="tc mono muted">{emitido}</td>'
            f'</tr>'
        )
    if not rows_nfse:
        rows_nfse = '<tr><td colspan="6" class="tc muted" style="padding:32px">Nenhuma NFS-e no período.</td></tr>'

    # ── Contas a Pagar ──
    cp_list  = contas_pagar or []
    cp_total = sum(float(c.get('valor',0)) for c in cp_list)
    cp_pago  = sum(float(c.get('valor',0)) for c in cp_list if c.get('status') == 'pago')
    cp_pend  = sum(float(c.get('valor',0)) for c in cp_list if c.get('status') == 'pendente')

    _ST_PAGAR_PRINT = {
        'pendente':  ('pend','#b45309','#fffbeb','#fde68a'),
        'pago':      ('ok',  '#15803d','#f0fdf4','#bbf7d0'),
        'cancelado': ('can', '#6b7280','#f9fafb','#e5e7eb'),
    }
    rows_pagar = ''
    for c in cp_list:
        st = c.get('status','pendente')
        cls_, col, bg, bord = _ST_PAGAR_PRINT.get(st, ('can','#6b7280','#f9fafb','#e5e7eb'))
        cat  = c.get('categoria_nome') or '—'
        obra = c.get('obra_nome') or '—'
        rows_pagar += (
            f'<tr>'
            f'<td>{c.get("descricao","—")}</td>'
            f'<td><span class="cat-tag">{cat}</span></td>'
            f'<td class="muted" style="font-size:11px">{obra}</td>'
            f'<td class="tr mono" style="color:#dc2626">{_fmt_brl(c.get("valor",0))}</td>'
            f'<td class="tc mono muted">{c.get("data_venc","—")}</td>'
            f'<td class="tc mono muted">{c.get("data_pag","—") if st=="pago" else "—"}</td>'
            f'<td class="tc"><span class="badge" style="color:{col};background:{bg};border-color:{bord}">{st.upper()}</span></td>'
            f'</tr>'
        )
    if not rows_pagar:
        rows_pagar = '<tr><td colspan="7" class="tc muted" style="padding:24px">Nenhuma conta a pagar.</td></tr>'

    # ── Contas a Receber ──
    cr_list  = contas_receber or []
    cr_total = sum(float(c.get('valor_total',0)) for c in cr_list)
    cr_receb = sum(float(c.get('valor_pago',0))  for c in cr_list)
    cr_pend  = round(cr_total - cr_receb, 2)

    _ST_RECEBER_PRINT = {
        'aberto':    ('#1d4ed8','#eff6ff','#bfdbfe'),
        'parcial':   ('#b45309','#fffbeb','#fde68a'),
        'quitado':   ('#15803d','#f0fdf4','#bbf7d0'),
        'cancelado': ('#6b7280','#f9fafb','#e5e7eb'),
    }
    rows_receber = ''
    for c in cr_list:
        st  = c.get('status','aberto')
        col, bg, bord = _ST_RECEBER_PRINT.get(st, ('#6b7280','#f9fafb','#e5e7eb'))
        vt  = float(c.get('valor_total',0))
        vp  = float(c.get('valor_pago',0))
        pct = int(min(100, vp/vt*100)) if vt else 0
        parc = c.get('parcelas',1)
        npag = len(c.get('pagamentos',[]))
        obra = c.get('obra_nome') or '—'
        rows_receber += (
            f'<tr>'
            f'<td>{c.get("descricao","—")}</td>'
            f'<td>{c.get("cliente_nome","—")}</td>'
            f'<td class="muted" style="font-size:11px">{obra}</td>'
            f'<td class="tr mono">{_fmt_brl(vt)}</td>'
            f'<td class="tr mono" style="color:#15803d">{_fmt_brl(vp)}</td>'
            f'<td class="tc" style="min-width:90px">'
            f'<div style="height:5px;border-radius:3px;background:#e5e7eb;overflow:hidden;margin-bottom:3px">'
            f'<div style="width:{pct}%;height:100%;background:{col}"></div></div>'
            f'<span style="font:10px monospace;color:#6b7280">{npag}/{parc}</span>'
            f'</td>'
            f'<td class="tc mono muted">{c.get("data_venc","—")}</td>'
            f'<td class="tc"><span class="badge" style="color:{col};background:{bg};border-color:{bord}">{st.upper()}</span></td>'
            f'</tr>'
        )
    if not rows_receber:
        rows_receber = '<tr><td colspan="8" class="tc muted" style="padding:24px">Nenhuma conta a receber.</td></tr>'

    # ── Saldo ──
    saldo = round(cr_receb - cp_pago, 2)
    saldo_col = '#15803d' if saldo >= 0 else '#dc2626'

    return f'''<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<title>Relatório Financeiro — {empresa}</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif; font-size: 13px; color: #1a1a1a; background: #f8f8f8; }}
  .wrap {{ max-width: 980px; margin: 0 auto; padding: 32px 24px; }}
  .rpt-header {{ display: flex; justify-content: space-between; align-items: flex-start; border-bottom: 3px solid #15803d; padding-bottom: 18px; margin-bottom: 24px; }}
  .rpt-empresa {{ font-size: 20px; font-weight: 700; color: #15803d; }}
  .rpt-cnpj    {{ font-size: 11px; color: #555; margin-top: 3px; font-family: monospace; }}
  .rpt-meta    {{ text-align: right; font-size: 11px; color: #555; line-height: 1.8; }}
  .rpt-periodo {{ font-weight: 600; color: #1a1a1a; font-size: 13px; }}
  .summary {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin-bottom: 24px; }}
  .summary-2 {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 12px; margin-bottom: 12px; }}
  .s-card {{ background: #fff; border: 1px solid #e5e7eb; border-radius: 10px; padding: 14px 16px; }}
  .s-card .lbl {{ font-size: 9px; font-weight: 600; letter-spacing: .1em; text-transform: uppercase; color: #6b7280; margin-bottom: 6px; }}
  .s-card .val {{ font-size: 18px; font-weight: 700; }}
  .s-card .sub {{ font-size: 10px; color: #9ca3af; margin-top: 2px; }}
  .c-green {{ color: #15803d; }} .c-blue {{ color: #1d4ed8; }} .c-amber {{ color: #b45309; }} .c-red {{ color: #dc2626; }}
  .rpt-section-label {{
    font-size: 11px; font-weight: 700; letter-spacing: .1em; text-transform: uppercase;
    color: #374151; margin: 24px 0 10px; display: flex; align-items: center; gap: 8px;
  }}
  .rpt-section-label::after {{ content:''; flex:1; height:1px; background:#e5e7eb; }}
  .rpt-card {{ background: #fff; border: 1px solid #e5e7eb; border-radius: 12px; overflow: hidden; margin-bottom: 20px; }}
  .rpt-card-hdr {{ padding: 12px 18px; border-bottom: 1px solid #e5e7eb; font-size: 11px; font-weight: 600; letter-spacing: .08em; text-transform: uppercase; color: #374151; display: flex; justify-content: space-between; align-items: center; }}
  table {{ width: 100%; border-collapse: collapse; }}
  thead th {{ padding: 10px 14px; text-align: left; font-size: 10px; font-weight: 600; letter-spacing: .07em; text-transform: uppercase; color: #6b7280; background: #f9fafb; border-bottom: 1px solid #e5e7eb; }}
  tbody tr {{ border-bottom: 1px solid #f3f4f6; }} tbody tr:last-child {{ border-bottom: none; }} tbody tr:hover {{ background: #fafafa; }}
  tbody td {{ padding: 9px 14px; vertical-align: middle; }}
  .desc {{ font-size: 11px; color: #6b7280; margin-top: 2px; }} .mono {{ font-family: 'DM Mono','Courier New',monospace; }}
  .muted {{ color: #6b7280; }} .tc {{ text-align: center; }} .tr {{ text-align: right; }}
  .num-col {{ font-weight: 600; color: #15803d; font-family: monospace; }}
  .badge {{ display: inline-block; padding: 2px 8px; border-radius: 4px; font-size: 9px; font-weight: 700; letter-spacing: .08em; border: 1px solid; }}
  .badge.ok  {{ color: #15803d; background: #f0fdf4; border-color: #bbf7d0; }}
  .badge.hom {{ color: #b45309; background: #fffbeb; border-color: #fde68a; }}
  .badge.err {{ color: #dc2626; background: #fef2f2; border-color: #fecaca; }}
  .cat-tag {{ font-size: 10px; padding: 1px 6px; border-radius: 3px; background: #f3f4f6; color: #374151; }}
  .totals-row td {{ padding: 10px 14px; font-weight: 700; background: #f9fafb; border-top: 2px solid #e5e7eb; }}
  .rpt-footer {{ font-size: 10px; color: #9ca3af; text-align: center; border-top: 1px solid #e5e7eb; padding-top: 14px; margin-top: 8px; }}
  .print-bar {{ position: fixed; top: 0; left: 0; right: 0; background: #15803d; color: #fff; padding: 10px 24px; display: flex; align-items: center; gap: 16px; font-size: 13px; font-weight: 500; z-index: 99; box-shadow: 0 2px 8px rgba(0,0,0,.15); }}
  .print-bar button {{ background: #fff; color: #15803d; border: none; border-radius: 6px; padding: 6px 18px; font-size: 12px; font-weight: 700; cursor: pointer; }}
  .print-bar .close-btn {{ background: rgba(255,255,255,.15); color: #fff; margin-left: auto; }}
  @media print {{ .print-bar {{ display: none !important; }} body {{ background: #fff; }} .wrap {{ padding: 0; margin-top: 0 !important; }} }}
  @media (max-width: 700px) {{ .summary,.summary-2 {{ grid-template-columns: repeat(2, 1fr); }} }}
</style>
</head>
<body>
<div class="print-bar">
  <span>📄 Relatório Financeiro — {empresa}</span>
  <button onclick="window.print()">🖨 Imprimir / Salvar PDF</button>
  <button class="close-btn" onclick="window.close()">✕ Fechar</button>
</div>

<div class="wrap" style="margin-top:52px">

  <div class="rpt-header">
    <div>
      <div class="rpt-empresa">{empresa}</div>
      <div class="rpt-cnpj">CNPJ: {cnpj_fmt}</div>
    </div>
    <div class="rpt-meta">
      <div class="rpt-periodo">{periodo_label}</div>
      <div>Gerado em {now_str}</div>
    </div>
  </div>

  <!-- Resumo geral -->
  <div class="summary">
    <div class="s-card">
      <div class="lbl">NFS-e emitidas</div>
      <div class="val c-green">{_fmt_brl(total)}</div>
      <div class="sub">{len(entries)} nota(s) · ISS {_fmt_brl(total_iss)}</div>
    </div>
    <div class="s-card">
      <div class="lbl">A Receber</div>
      <div class="val c-blue">{_fmt_brl(cr_pend)}</div>
      <div class="sub">de {_fmt_brl(cr_total)} total</div>
    </div>
    <div class="s-card">
      <div class="lbl">A Pagar</div>
      <div class="val c-amber">{_fmt_brl(cp_pend)}</div>
      <div class="sub">de {_fmt_brl(cp_total)} total</div>
    </div>
    <div class="s-card">
      <div class="lbl">Saldo (Recebido − Pago)</div>
      <div class="val" style="color:{saldo_col}">{_fmt_brl(saldo)}</div>
      <div class="sub">recebido {_fmt_brl(cr_receb)} · pago {_fmt_brl(cp_pago)}</div>
    </div>
  </div>

  <!-- NFS-e -->
  <div class="rpt-section-label">Notas Fiscais de Serviço</div>
  <div class="rpt-card">
    <div class="rpt-card-hdr">
      <span>NFS-e — {periodo_label}</span>
      <span style="font-weight:400;color:#9ca3af">{len(entries)} nota(s) · Liq. {_fmt_brl(total_liq)}</span>
    </div>
    <table>
      <thead><tr>
        <th style="width:60px">Nº</th><th>Tomador / Descrição</th>
        <th style="width:120px;text-align:right">Valor</th>
        <th style="width:100px;text-align:right">ISS</th>
        <th style="width:90px;text-align:center">Status</th>
        <th style="width:90px;text-align:center">Emitida</th>
      </tr></thead>
      <tbody>
        {rows_nfse}
        <tr class="totals-row">
          <td colspan="2" class="muted" style="font-size:11px">Total — {len(entries)} nota(s)</td>
          <td class="tr mono c-green">{_fmt_brl(total)}</td>
          <td class="tr mono c-amber">{_fmt_brl(total_iss)}</td>
          <td colspan="2"></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Contas a Pagar -->
  <div class="rpt-section-label">Contas a Pagar</div>
  <div class="rpt-card">
    <div class="rpt-card-hdr">
      <span>Contas a Pagar</span>
      <span style="font-weight:400;color:#9ca3af">Total {_fmt_brl(cp_total)} · Pago {_fmt_brl(cp_pago)} · Pendente {_fmt_brl(cp_pend)}</span>
    </div>
    <table>
      <thead><tr>
        <th>Descrição</th><th style="width:110px">Categoria</th>
        <th style="width:130px">Obra</th>
        <th style="width:110px;text-align:right">Valor</th>
        <th style="width:90px;text-align:center">Vencimento</th>
        <th style="width:90px;text-align:center">Pagamento</th>
        <th style="width:80px;text-align:center">Status</th>
      </tr></thead>
      <tbody>
        {rows_pagar}
        <tr class="totals-row">
          <td colspan="3" class="muted" style="font-size:11px">{len(cp_list)} conta(s)</td>
          <td class="tr mono" style="color:#dc2626">{_fmt_brl(cp_total)}</td>
          <td colspan="3"></td>
        </tr>
      </tbody>
    </table>
  </div>

  <!-- Contas a Receber -->
  <div class="rpt-section-label">Contas a Receber</div>
  <div class="rpt-card">
    <div class="rpt-card-hdr">
      <span>Contas a Receber</span>
      <span style="font-weight:400;color:#9ca3af">Total {_fmt_brl(cr_total)} · Recebido {_fmt_brl(cr_receb)} · Pendente {_fmt_brl(cr_pend)}</span>
    </div>
    <table>
      <thead><tr>
        <th>Descrição</th><th style="width:130px">Cliente</th>
        <th style="width:130px">Obra</th>
        <th style="width:110px;text-align:right">Total</th>
        <th style="width:110px;text-align:right">Recebido</th>
        <th style="width:80px;text-align:center">Parcelas</th>
        <th style="width:90px;text-align:center">Vencimento</th>
        <th style="width:80px;text-align:center">Status</th>
      </tr></thead>
      <tbody>
        {rows_receber}
        <tr class="totals-row">
          <td colspan="3" class="muted" style="font-size:11px">{len(cr_list)} conta(s)</td>
          <td class="tr mono">{_fmt_brl(cr_total)}</td>
          <td class="tr mono c-green">{_fmt_brl(cr_receb)}</td>
          <td colspan="3"></td>
        </tr>
      </tbody>
    </table>
  </div>

  <div class="rpt-footer">
    DMC Topografia · Sistema de Gestão · Relatório gerado automaticamente em {now_str}
  </div>

</div>
</body>
</html>'''


_MESES_PT = {
    '01':'Janeiro','02':'Fevereiro','03':'Março','04':'Abril',
    '05':'Maio','06':'Junho','07':'Julho','08':'Agosto',
    '09':'Setembro','10':'Outubro','11':'Novembro','12':'Dezembro',
}
_MESES_ABBR = {
    '01':'Jan','02':'Fev','03':'Mar','04':'Abr','05':'Mai','06':'Jun',
    '07':'Jul','08':'Ago','09':'Set','10':'Out','11':'Nov','12':'Dez',
}


def _gerar_word_doc(tipo: str, mes: str, ano: str, cfg: dict) -> bytes:
    """Gera relatório financeiro .docx e retorna os bytes."""
    import io
    from collections import defaultdict
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from services.financeiro import (
        load_contas_pagar, load_contas_receber, load_categorias_pagar,
    )
    from services.nfse import list_nfse

    all_pagar   = load_contas_pagar()
    all_receber = load_contas_receber()
    all_nfse    = list_nfse()
    all_cats    = load_categorias_pagar()

    empresa  = cfg.get('razao_social') or 'Empresa'
    cnpj_raw = cfg.get('cnpj') or ''
    cnpj_fmt = (
        f'{cnpj_raw[:2]}.{cnpj_raw[2:5]}.{cnpj_raw[5:8]}/{cnpj_raw[8:12]}-{cnpj_raw[12:]}'
        if len(cnpj_raw) == 14 else cnpj_raw
    )
    aliq     = float(cfg.get('aliquota_iss', '5.00') or '5.00')
    hoje_str = datetime.now().strftime('%d/%m/%Y %H:%M')
    mes_nome = _MESES_PT.get(mes, mes)
    mes_ano_str = f"{mes_nome} / {ano}"
    mes_ano  = f"{ano}-{mes}"

    def _d2m(s: str) -> str:
        s = (s or '').strip()
        if len(s) >= 7 and s[4:5] == '-':
            return s[:7]
        if len(s) >= 7 and s[2:3] == '/':
            return f"{s[6:10]}-{s[3:5]}"
        return ''

    def _brl(v) -> str:
        try:
            f = float(v or 0)
        except (ValueError, TypeError):
            f = 0.0
        return f'R$ {f:,.2f}'.replace(',','X').replace('.', ',').replace('X','.')

    def _filtro_mes(items, campo, fb='data_criacao'):
        out = []
        for i in items:
            v = _d2m(i.get(campo) or '')
            if not v:
                v = _d2m(i.get(fb) or '')
            if v == mes_ano:
                out.append(i)
        return out

    pagar_mes   = _filtro_mes(all_pagar,   'data_venc')
    receber_mes = _filtro_mes(all_receber, 'data_venc')
    nfse_mes    = [e for e in all_nfse if _d2m(e.get('emitido_em','') or '') == mes_ano]

    # ── Document setup ────────────────────────────────────────────────
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(0.9)
        sec.right_margin  = Inches(0.9)

    R = WD_ALIGN_PARAGRAPH.RIGHT
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    def _rgb(hex_c: str):
        h = hex_c.lstrip('#')
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    def _set_cell_bg(cell, hex_c: str):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_c.lstrip('#'))
        tcPr.append(shd)

    def _doc_header(title: str, periodo: str):
        h = doc.add_heading(title, level=0)
        h.alignment = C
        if h.runs:
            h.runs[0].font.color.rgb = _rgb('#15803d')
            h.runs[0].font.size = Pt(20)
        p = doc.add_paragraph()
        p.alignment = C
        r = p.add_run(f"{empresa}  |  CNPJ: {cnpj_fmt}")
        r.font.size = Pt(10)
        r.font.color.rgb = _rgb('#6b7280')
        p2 = doc.add_paragraph()
        p2.alignment = C
        r2 = p2.add_run(f"Período: {periodo}  ·  Gerado em {hoje_str}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = _rgb('#9ca3af')
        doc.add_paragraph()

    def _section_heading(title: str, color: str = '#374151'):
        h = doc.add_heading(title, level=1)
        if h.runs:
            h.runs[0].font.color.rgb = _rgb(color)
            h.runs[0].font.size = Pt(12)

    def _kpi_table(items):
        cols = min(len(items), 4)
        rows = (len(items) + cols - 1) // cols
        tbl = doc.add_table(rows=rows, cols=cols)
        tbl.style = 'Table Grid'
        for idx, (lbl, val, col) in enumerate(items):
            cell = tbl.cell(idx // cols, idx % cols)
            _set_cell_bg(cell, 'F9FAFB')
            p = cell.paragraphs[0]
            p.clear()
            rl = p.add_run(lbl.upper() + '\n')
            rl.font.size = Pt(8)
            rl.font.color.rgb = _rgb('#6b7280')
            rv = p.add_run(val)
            rv.bold = True
            rv.font.size = Pt(14)
            rv.font.color.rgb = _rgb(col)
        doc.add_paragraph()

    def _data_table(headers, rows_data, totals_row=None):
        ncols = len(headers)
        nrows = 1 + len(rows_data) + (1 if totals_row else 0)
        tbl = doc.add_table(rows=nrows, cols=ncols)
        tbl.style = 'Table Grid'
        # Header
        for ci, h in enumerate(headers):
            cell = tbl.rows[0].cells[ci]
            _set_cell_bg(cell, 'F3F4F6')
            p = cell.paragraphs[0]
            p.alignment = C
            run = p.add_run(h.upper())
            run.bold = True
            run.font.size = Pt(8)
            run.font.color.rgb = _rgb('#6b7280')
        # Data
        for ri, row_data in enumerate(rows_data):
            for ci, cd in enumerate(row_data):
                cell = tbl.rows[ri + 1].cells[ci]
                if isinstance(cd, tuple):
                    text  = str(cd[0])
                    color = cd[1] if len(cd) > 1 else None
                    align = cd[2] if len(cd) > 2 else L
                    bold  = cd[3] if len(cd) > 3 else False
                else:
                    text, color, align, bold = str(cd), None, L, False
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.add_run(text)
                run.font.size = Pt(9)
                run.bold = bold
                if color:
                    run.font.color.rgb = _rgb(color)
        # Totals
        if totals_row:
            for ci, cd in enumerate(totals_row):
                cell = tbl.rows[-1].cells[ci]
                _set_cell_bg(cell, 'F9FAFB')
                if isinstance(cd, tuple):
                    text  = str(cd[0])
                    color = cd[1] if len(cd) > 1 else None
                    align = cd[2] if len(cd) > 2 else L
                else:
                    text, color, align = str(cd), None, L
                p = cell.paragraphs[0]
                p.alignment = align
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(9)
                if color:
                    run.font.color.rgb = _rgb(color)
        doc.add_paragraph()

    # ── a. Balanço Geral Mensal ───────────────────────────────────────
    if tipo == 'balanco_mes':
        _doc_header('Balanço Geral Mensal', mes_ano_str)
        nfse_t  = sum(float(e.get('valor',0) or 0) for e in nfse_mes)
        cp_t    = sum(float(c.get('valor',0) or 0) for c in pagar_mes)
        cp_pago = sum(float(c.get('valor',0) or 0) for c in pagar_mes if c.get('status')=='pago')
        cr_t    = sum(float(c.get('valor_total',0) or 0) for c in receber_mes)
        cr_r    = sum(float(c.get('valor_pago',0) or 0) for c in receber_mes)
        saldo   = round(cr_r - cp_pago, 2)
        _kpi_table([
            ('NFS-e Emitidas',     _brl(nfse_t),  '#15803d'),
            ('Receitas Recebidas', _brl(cr_r),     '#1d4ed8'),
            ('Despesas Pagas',     _brl(cp_pago),  '#dc2626'),
            ('Saldo',              _brl(saldo),    '#15803d' if saldo >= 0 else '#dc2626'),
        ])
        _section_heading('Notas Fiscais de Serviço', '#15803d')
        _data_table(
            ['Nº','Tomador','Descrição','Valor','ISS','Status','Emitida'],
            [[
                f'#{e.get("numero","—")}',
                e.get('tomador','—'),
                (e.get('descricao','—') or '')[:50],
                (_brl(float(e.get('valor',0) or 0)), '#15803d', R),
                (_brl(round(float(e.get('valor',0) or 0)*aliq/100,2)), '#6b7280', R),
                ('EMITIDA' if (e.get('sucesso') and e.get('ambiente')=='producao')
                 else ('HOMO' if e.get('sucesso') else 'ERRO'),
                 '#15803d' if (e.get('sucesso') and e.get('ambiente')=='producao')
                 else ('#b45309' if e.get('sucesso') else '#dc2626'), C),
                ((e.get('emitido_em','') or '')[:10], None, C),
            ] for e in nfse_mes],
            totals_row=[(f'Total: {len(nfse_mes)} nota(s)','#6b7280',L),'','',
                        (_brl(nfse_t),'#15803d',R),'','',''],
        )
        _section_heading('Contas a Pagar', '#dc2626')
        _data_table(
            ['Descrição','Categoria','Obra','Valor','Vencimento','Pagamento','Status'],
            [[
                c.get('descricao','—'),
                c.get('categoria_nome','—'),
                c.get('obra_nome','—'),
                (_brl(c.get('valor',0)), '#dc2626', R),
                (c.get('data_venc','—'), None, C),
                (c.get('data_pag','—') if c.get('status')=='pago' else '—', None, C),
                (c.get('status','pendente').upper(),
                 {'pago':'#15803d','pendente':'#b45309','cancelado':'#6b7280'}.get(c.get('status','pendente'),'#6b7280'), C),
            ] for c in pagar_mes],
            totals_row=[(f'{len(pagar_mes)} conta(s)','#6b7280',L),'','',
                        (_brl(cp_t),'#dc2626',R),'','',''],
        )
        _section_heading('Contas a Receber', '#1d4ed8')
        _data_table(
            ['Descrição','Cliente','Obra','Total','Recebido','Vencimento','Status'],
            [[
                c.get('descricao','—'),
                c.get('cliente_nome','—'),
                c.get('obra_nome','—'),
                (_brl(c.get('valor_total',0)), '#1d4ed8', R),
                (_brl(c.get('valor_pago',0)), '#15803d', R),
                (c.get('data_venc','—'), None, C),
                (c.get('status','aberto').upper(),
                 {'quitado':'#15803d','parcial':'#b45309','aberto':'#1d4ed8','cancelado':'#6b7280'}.get(c.get('status','aberto'),'#6b7280'), C),
            ] for c in receber_mes],
            totals_row=[(f'{len(receber_mes)} conta(s)','#6b7280',L),'','',
                        (_brl(cr_t),'#1d4ed8',R),(_brl(cr_r),'#15803d',R),'',''],
        )

    # ── b. Balanço Geral Anual ────────────────────────────────────────
    elif tipo == 'balanco_ano':
        _doc_header('Balanço Geral Anual', ano)
        _section_heading('Resumo por Mês', '#374151')
        meses_list = ['01','02','03','04','05','06','07','08','09','10','11','12']
        ano_rows = []
        tt_nfse = tt_cp = tt_cr = tt_receb = tt_pago = 0.0
        for m in meses_list:
            ma = f"{ano}-{m}"
            ne  = [e for e in all_nfse    if _d2m(e.get('emitido_em','') or '') == ma]
            cp  = [c for c in all_pagar   if (_d2m(c.get('data_venc','') or '') or _d2m(c.get('data_criacao','') or '')) == ma]
            cr  = [c for c in all_receber if (_d2m(c.get('data_venc','') or '') or _d2m(c.get('data_criacao','') or '')) == ma]
            nv  = sum(float(e.get('valor',0) or 0) for e in ne)
            cpv = sum(float(c.get('valor',0) or 0) for c in cp)
            crv = sum(float(c.get('valor_total',0) or 0) for c in cr)
            crr = sum(float(c.get('valor_pago',0)  or 0) for c in cr)
            cpp = sum(float(c.get('valor',0) or 0) for c in cp if c.get('status')=='pago')
            sl  = round(crr - cpp, 2)
            tt_nfse += nv; tt_cp += cpv; tt_cr += crv; tt_receb += crr; tt_pago += cpp
            ano_rows.append([
                (_MESES_ABBR.get(m,m)+f'/{ano[-2:]}', None, L, True),
                (f'{len(ne)}', '#6b7280', C),
                (_brl(nv), '#15803d', R),
                (_brl(cpv), '#dc2626', R),
                (_brl(crv), '#1d4ed8', R),
                (_brl(crr), '#15803d', R),
                (_brl(sl), '#15803d' if sl >= 0 else '#dc2626', R, True),
            ])
        tt_sl = round(tt_receb - tt_pago, 2)
        _data_table(
            ['Mês','NFS-e','Fat. NFS-e','Total Pagar','Total Receber','Recebido','Saldo'],
            ano_rows,
            totals_row=[
                ('TOTAL ANUAL','#374151',L),'',
                (_brl(tt_nfse),'#15803d',R),
                (_brl(tt_cp),'#dc2626',R),
                (_brl(tt_cr),'#1d4ed8',R),
                (_brl(tt_receb),'#15803d',R),
                (_brl(tt_sl),'#15803d' if tt_sl>=0 else '#dc2626',R),
            ],
        )

    # ── c. Contas a Pagar Mensal ──────────────────────────────────────
    elif tipo == 'pagar_mes':
        _doc_header('Contas a Pagar — Mensal', mes_ano_str)
        cp_t  = sum(float(c.get('valor',0) or 0) for c in pagar_mes)
        cp_pg = sum(float(c.get('valor',0) or 0) for c in pagar_mes if c.get('status')=='pago')
        cp_pe = sum(float(c.get('valor',0) or 0) for c in pagar_mes if c.get('status')=='pendente')
        cp_ca = sum(float(c.get('valor',0) or 0) for c in pagar_mes if c.get('status')=='cancelado')
        _kpi_table([
            ('Total',     _brl(cp_t),  '#374151'),
            ('Pago',      _brl(cp_pg), '#15803d'),
            ('Pendente',  _brl(cp_pe), '#b45309'),
            ('Cancelado', _brl(cp_ca), '#6b7280'),
        ])
        _section_heading('Lançamentos', '#dc2626')
        _data_table(
            ['Descrição','Categoria','Obra','Valor','Vencimento','Pagamento','Obs.','Status'],
            [[
                c.get('descricao','—'),
                c.get('categoria_nome','—'),
                c.get('obra_nome','—'),
                (_brl(c.get('valor',0)), '#dc2626', R),
                (c.get('data_venc','—'), None, C),
                (c.get('data_pag','—') if c.get('status')=='pago' else '—', None, C),
                ((c.get('observacao','') or '')[:40], '#6b7280', L),
                (c.get('status','pendente').upper(),
                 {'pago':'#15803d','pendente':'#b45309','cancelado':'#6b7280'}.get(c.get('status','pendente'),'#6b7280'), C),
            ] for c in pagar_mes],
            totals_row=[(f'{len(pagar_mes)} conta(s)','#6b7280',L),'','',
                        (_brl(cp_t),'#dc2626',R),'','','',''],
        )

    # ── d. Contas a Receber Mensal ────────────────────────────────────
    elif tipo == 'receber_mes':
        _doc_header('Contas a Receber — Mensal', mes_ano_str)
        cr_t = sum(float(c.get('valor_total',0) or 0) for c in receber_mes)
        cr_r = sum(float(c.get('valor_pago',0)  or 0) for c in receber_mes)
        cr_p = round(cr_t - cr_r, 2)
        _kpi_table([
            ('Total a Receber', _brl(cr_t), '#1d4ed8'),
            ('Recebido',        _brl(cr_r), '#15803d'),
            ('Pendente',        _brl(cr_p), '#b45309'),
        ])
        _section_heading('Lançamentos', '#1d4ed8')
        _data_table(
            ['Descrição','Cliente','CPF','Obra','Total','Recebido','Parcelas','Vencimento','Status'],
            [[
                c.get('descricao','—'),
                c.get('cliente_nome','—'),
                c.get('cliente_cpf','—'),
                c.get('obra_nome','—'),
                (_brl(c.get('valor_total',0)), '#1d4ed8', R),
                (_brl(c.get('valor_pago',0)), '#15803d', R),
                (f'{len(c.get("pagamentos",[]))}/{c.get("parcelas",1)}', '#6b7280', C),
                (c.get('data_venc','—'), None, C),
                (c.get('status','aberto').upper(),
                 {'quitado':'#15803d','parcial':'#b45309','aberto':'#1d4ed8','cancelado':'#6b7280'}.get(c.get('status','aberto'),'#6b7280'), C),
            ] for c in receber_mes],
            totals_row=[(f'{len(receber_mes)} conta(s)','#6b7280',L),'','','',
                        (_brl(cr_t),'#1d4ed8',R),(_brl(cr_r),'#15803d',R),'','',''],
        )

    # ── e. Por Categoria Mensal ───────────────────────────────────────
    elif tipo == 'categoria_mes':
        _doc_header('Relatório por Categoria Mensal', mes_ano_str)
        cats_map = {c['id']: c for c in all_cats}
        by_cat: dict = defaultdict(list)
        sem_cat = []
        for c in pagar_mes:
            cid = c.get('categoria_id','')
            if cid and cid in cats_map:
                by_cat[cid].append(c)
            else:
                sem_cat.append(c)
        grand = 0.0
        for cid, contas in sorted(by_cat.items(), key=lambda x: cats_map[x[0]].get('nome','').lower()):
            cat_nome = cats_map[cid].get('nome','—')
            sub = sum(float(c.get('valor',0) or 0) for c in contas)
            grand += sub
            _section_heading(f'Categoria: {cat_nome}', '#374151')
            _data_table(
                ['Descrição','Obra','Valor','Vencimento','Status'],
                [[
                    c.get('descricao','—'), c.get('obra_nome','—'),
                    (_brl(c.get('valor',0)), '#dc2626', R),
                    (c.get('data_venc','—'), None, C),
                    (c.get('status','pendente').upper(),
                     {'pago':'#15803d','pendente':'#b45309','cancelado':'#6b7280'}.get(c.get('status','pendente'),'#6b7280'), C),
                ] for c in contas],
                totals_row=[(f'Subtotal — {cat_nome}','#374151',L),'',
                            (_brl(sub),'#dc2626',R),'',''],
            )
        if sem_cat:
            sub_sc = sum(float(c.get('valor',0) or 0) for c in sem_cat)
            grand += sub_sc
            _section_heading('Sem Categoria', '#6b7280')
            _data_table(
                ['Descrição','Obra','Valor','Vencimento','Status'],
                [[
                    c.get('descricao','—'), c.get('obra_nome','—'),
                    (_brl(c.get('valor',0)), '#dc2626', R),
                    (c.get('data_venc','—'), None, C),
                    (c.get('status','pendente').upper(),
                     {'pago':'#15803d','pendente':'#b45309','cancelado':'#6b7280'}.get(c.get('status','pendente'),'#6b7280'), C),
                ] for c in sem_cat],
                totals_row=[('Subtotal — Sem Categoria','#6b7280',L),'',
                            (_brl(sub_sc),'#dc2626',R),'',''],
            )
        p = doc.add_paragraph()
        p.alignment = R
        run = p.add_run(f'TOTAL GERAL: {_brl(grand)}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = _rgb('#dc2626')

    # ── Footer ────────────────────────────────────────────────────────
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = C
    fr = fp.add_run(f'DMC Topografia · Sistema de Gestão · {hoje_str}')
    fr.font.size = Pt(8)
    fr.font.color.rgb = _rgb('#9ca3af')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def relatorio_financeiro_dialog() -> None:
    cfg = load_config()
    hoje = datetime.now()

    _TIPOS = [
        ('balanco_mes',   'a. Balanço Geral por Mês'),
        ('balanco_ano',   'b. Balanço Geral por Ano'),
        ('pagar_mes',     'c. Balanço Contas a Pagar Mensal'),
        ('receber_mes',   'd. Balanço Contas a Receber Mensal'),
        ('categoria_mes', 'e. Relatório por Categoria Mensal'),
    ]
    _TIPOS_SEM_MES = {'balanco_ano'}

    anos_disponiveis = [str(y) for y in range(hoje.year, hoje.year - 6, -1)]
    meses_disponiveis = [
        ('01','Janeiro'),('02','Fevereiro'),('03','Março'),('04','Abril'),
        ('05','Maio'),('06','Junho'),('07','Julho'),('08','Agosto'),
        ('09','Setembro'),('10','Outubro'),('11','Novembro'),('12','Dezembro'),
    ]
    mes_cur = f'{hoje.month:02d}'
    ano_cur = str(hoje.year)

    with ui.dialog() as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(560px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        # ── Cabeçalho ──────────────────────────────────────────────────
        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(96,165,250,.08);border:1px solid rgba(96,165,250,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#60A5FA">bar_chart</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Relatório Financeiro</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">'
                        'Gera arquivo Word (.docx) para download</div>')

        # ── Corpo ──────────────────────────────────────────────────────
        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0;'
                                     'display:flex;flex-direction:column;gap:16px'):

            # Tipo de relatório
            _tipo_id = f'rfi-tipo-{id(dlg)}'
            ui.html('<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
                    'letter-spacing:.14em;text-transform:uppercase;margin-bottom:6px">Tipo de Relatório</div>')
            _tipo_opts = ''.join(
                f'<option value="{k}"{"  selected" if k == "balanco_mes" else ""}>{v}</option>'
                for k, v in _TIPOS
            )
            ui.html(f'<select id="{_tipo_id}" class="dmc-input" style="cursor:pointer">{_tipo_opts}</select>')

            # Período
            ui.html('<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
                    'letter-spacing:.14em;text-transform:uppercase;margin-bottom:6px">Período</div>')

            _mes_box_id = f'rfi-mesbox-{id(dlg)}'
            _mes_id     = f'rfi-mes-{id(dlg)}'
            _ano_id     = f'rfi-ano-{id(dlg)}'

            _mes_opts = ''.join(
                f'<option value="{k}"{"  selected" if k == mes_cur else ""}>{v}</option>'
                for k, v in meses_disponiveis
            )
            _ano_opts = ''.join(
                f'<option value="{a}"{"  selected" if a == ano_cur else ""}>{a}</option>'
                for a in anos_disponiveis
            )
            ui.html(
                f'<div style="display:grid;grid-template-columns:1fr 120px;gap:10px">'
                f'<div id="{_mes_box_id}">'
                f'<select id="{_mes_id}" class="dmc-input" style="cursor:pointer">{_mes_opts}</select>'
                f'</div>'
                f'<select id="{_ano_id}" class="dmc-input" style="cursor:pointer">{_ano_opts}</select>'
                f'</div>'
            )

            # Tipos de relatório disponíveis (descrição)
            ui.html(
                '<div style="background:rgba(96,165,250,.06);border:1px solid rgba(96,165,250,.18);'
                'border-radius:10px;padding:12px 14px">'
                '<div style="font:600 11px var(--dmc-mono);color:#60A5FA;margin-bottom:8px;'
                'text-transform:uppercase;letter-spacing:.08em">Tipos disponíveis</div>'
                '<div style="display:flex;flex-direction:column;gap:5px;font:12px var(--dmc-fm);color:var(--dmc-muted2)">'
                '<span><b style="color:var(--dmc-text)">a. Balanço Geral por Mês</b> — NFS-e, Pagar e Receber do mês</span>'
                '<span><b style="color:var(--dmc-text)">b. Balanço Geral por Ano</b> — tabela com os 12 meses do ano</span>'
                '<span><b style="color:var(--dmc-text)">c. Contas a Pagar Mensal</b> — todos os lançamentos do mês</span>'
                '<span><b style="color:var(--dmc-text)">d. Contas a Receber Mensal</b> — todos os recebimentos do mês</span>'
                '<span><b style="color:var(--dmc-text)">e. Por Categoria Mensal</b> — despesas agrupadas por categoria</span>'
                '</div></div>'
            )

            async def _setup_tipo():
                await ui.run_javascript(
                    f'(function(){{'
                    f'var sel=document.getElementById("{_tipo_id}");'
                    f'var mb=document.getElementById("{_mes_box_id}");'
                    f'var semMes={list(_TIPOS_SEM_MES)};'
                    f'function upd(){{if(mb)mb.style.display=semMes.indexOf(sel.value)>=0?"none":"block";}}'
                    f'if(sel)sel.addEventListener("change",upd);'
                    f'upd();'
                    f'}})()'
                )
            ui.timer(0.15, _setup_tipo, once=True)

        # ── Rodapé ─────────────────────────────────────────────────────
        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            _gerar_btn_ref: dict = {}

            async def _gerar():
                btn = _gerar_btn_ref.get('btn')
                if btn:
                    btn.props('loading=true disabled=true')
                try:
                    vals = await ui.run_javascript(
                        f'(function(){{'
                        f'var t=document.getElementById("{_tipo_id}");'
                        f'var m=document.getElementById("{_mes_id}");'
                        f'var a=document.getElementById("{_ano_id}");'
                        f'return{{tipo:t?t.value:"balanco_mes",'
                        f'mes:m?m.value:"{mes_cur}",'
                        f'ano:a?a.value:"{ano_cur}"}};'
                        f'}})()'
                    )
                    if not isinstance(vals, dict):
                        ui.notify('Erro ao ler formulário', color='negative')
                        return
                    tipo = vals.get('tipo', 'balanco_mes')
                    mes  = vals.get('mes', mes_cur)
                    ano  = vals.get('ano', ano_cur)

                    doc_bytes = _gerar_word_doc(tipo, mes, ano, cfg)
                    b64 = base64.b64encode(doc_bytes).decode()

                    tipo_nome = dict(_TIPOS).get(tipo, tipo).replace(' ','_')
                    mes_nome  = _MESES_PT.get(mes, mes)
                    if tipo == 'balanco_ano':
                        fname = f'Relatorio_{tipo_nome}_{ano}.docx'
                    else:
                        fname = f'Relatorio_{tipo_nome}_{mes_nome}_{ano}.docx'
                    fname = fname.replace(' ','_').replace('.','_').replace('___','_') + '.docx'
                    fname = fname.replace('.docx.docx', '.docx')

                    await ui.run_javascript(
                        f'(function(){{'
                        f'var a=document.createElement("a");'
                        f'a.href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}";'
                        f'a.download="{fname}";'
                        f'document.body.appendChild(a);a.click();document.body.removeChild(a);'
                        f'}})()'
                    )
                    ui.notify('Relatório gerado com sucesso!', color='positive', icon='check_circle')
                    dlg.close()
                except Exception as ex:
                    ui.notify(f'Erro ao gerar relatório: {ex}', color='negative')
                finally:
                    if btn:
                        btn.props('loading=false disabled=false')

            btn = ui.button('Gerar Word', icon='description', on_click=_gerar).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary').style(
                'background:rgba(96,165,250,.12);border-color:rgba(96,165,250,.35);color:#60A5FA'
            )
            _gerar_btn_ref['btn'] = btn

    dlg.open()


# ═══════════════════════════════════════════════════════════════════════════
# CONTAS A PAGAR
# ═══════════════════════════════════════════════════════════════════════════

_CAT_COLORS = [
    "#4ADE80","#FBBF24","#60A5FA","#F87171","#C4B5FD",
    "#FB923C","#34D399","#E879F9","#38BDF8","#A3E635",
]


def categorias_pagar_dialog(on_save=None) -> None:
    """CRUD de categorias de contas a pagar."""
    from services.financeiro import (
        load_categorias_pagar, add_categoria_pagar,
        update_categoria_pagar, delete_categoria_pagar,
    )

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(500px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(251,191,36,.08);border:1px solid rgba(251,191,36,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#FBBF24">label</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Categorias de Despesa</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Tipos de conta a pagar — gasolina, vale alimentação, etc.</div>')

        list_area = ui.element('div').style(
            'padding:16px 24px;overflow-y:auto;flex:1;min-height:0;'
            'display:flex;flex-direction:column;gap:8px'
        )

        def _redraw():
            list_area.clear()
            cats = load_categorias_pagar()
            with list_area:
                if not cats:
                    ui.html(
                        '<div style="text-align:center;padding:32px;'
                        'font:11px var(--dmc-mono);color:var(--dmc-muted2)">Nenhuma categoria cadastrada</div>'
                    )
                for cat in cats:
                    with ui.element('div').style(
                        'display:flex;align-items:center;gap:10px;'
                        'padding:10px 14px;background:rgba(255,255,255,.03);'
                        'border:1px solid var(--dmc-b1);border-radius:10px;'
                    ):
                        cor = cat.get('cor','#8BAA8B')
                        ui.html(
                            f'<div style="width:14px;height:14px;border-radius:4px;'
                            f'flex-shrink:0;background:{cor}"></div>'
                        )
                        ui.html(
                            f'<span style="font:500 13px var(--dmc-fm);'
                            f'color:var(--dmc-text);flex:1">{cat["nome"]}</span>'
                        )
                        cid = cat['id']
                        del_btn = ui.element('button').style(
                            'width:28px;height:28px;border-radius:6px;cursor:pointer;'
                            'background:transparent;border:1px solid var(--dmc-b1);'
                            'display:inline-flex;align-items:center;justify-content:center;'
                            'color:#F87171;border-color:rgba(248,113,113,.2)'
                        )
                        with del_btn:
                            ui.html('<span class="material-icons" style="font-size:14px">delete</span>')
                        del_btn.on('click', lambda c=cid: (_del(c)))

                def _del(cid):
                    delete_categoria_pagar(cid)
                    _redraw()
                    if on_save:
                        on_save()

        _redraw()

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;flex-direction:column;gap:10px;flex-shrink:0'
        ):
            ui.html('<label class="dmc-label">Nova categoria</label>')
            _new_nome_id  = f'cat-nome-{id(dlg)}'
            _new_cor_ref  = {'v': _CAT_COLORS[0]}
            with ui.element('div').style('display:flex;gap:8px;align-items:center'):
                ui.html(
                    f'<input id="{_new_nome_id}" class="dmc-input" placeholder="ex: Gasolina" '
                    f'style="flex:1">'
                )
                cor_box = ui.element('div').style(
                    f'width:34px;height:34px;border-radius:8px;flex-shrink:0;cursor:pointer;'
                    f'background:{_CAT_COLORS[0]};border:2px solid rgba(255,255,255,.12)'
                )
                with ui.element('div').style('display:flex;gap:6px;flex-wrap:wrap;max-width:200px'):
                    for c in _CAT_COLORS:
                        dot = ui.element('div').style(
                            f'width:18px;height:18px;border-radius:5px;cursor:pointer;'
                            f'background:{c};border:2px solid transparent'
                        )
                        def _pick(col=c):
                            _new_cor_ref['v'] = col
                            cor_box.style(
                                f'width:34px;height:34px;border-radius:8px;flex-shrink:0;cursor:pointer;'
                                f'background:{col};border:2px solid rgba(255,255,255,.12)'
                            )
                        dot.on('click', _pick)

            async def _add_cat():
                nome = await ui.run_javascript(
                    f'(document.getElementById("{_new_nome_id}")||{{}}).value||""'
                )
                nome = (nome or '').strip()
                if not nome:
                    return
                add_categoria_pagar(nome, _new_cor_ref['v'])
                await ui.run_javascript(
                    f'var e=document.getElementById("{_new_nome_id}");if(e)e.value="";'
                )
                _redraw()
                if on_save:
                    on_save()

            with ui.element('div').style('display:flex;justify-content:flex-end;gap:10px'):
                ui.button('Fechar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')
                ui.button('Adicionar', icon='add', on_click=_add_cat).props(
                    'unelevated no-caps'
                ).classes('dmc-btn dmc-btn-primary').style(
                    'background:rgba(251,191,36,.12);border-color:rgba(251,191,36,.35);color:#FBBF24'
                )

    dlg.open()


def nova_conta_pagar_dialog(conta: dict | None = None, on_save=None) -> None:
    """Cadastro / edição de conta a pagar."""
    from services.financeiro import (
        load_categorias_pagar, add_categoria_pagar,
        add_conta_pagar, update_conta_pagar,
    )
    from services.obras import load_obras

    editando = conta is not None
    cats  = load_categorias_pagar()
    obras = load_obras()
    obra_opts = [('', '— Nenhuma —')] + [(o['id'], o.get('cliente_nome','') + ' · ' + o.get('obra_log','')) for o in obras]
    cat_opts  = [('', '— Sem categoria —')] + [(c['id'], c['nome']) for c in cats]

    _st: dict = dict(conta) if conta else {}

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(560px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#F87171">arrow_circle_down</span></div>'
            )
            with ui.element('div'):
                ui.html(f'<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">{"Editar" if editando else "Nova"} Conta a Pagar</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Registre despesas e vincule a obras</div>')

        _iids = {}

        def _inp_id(key): return f'cp-{key}-{id(dlg)}'

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0;display:flex;flex-direction:column;gap:12px'):

            ui.html('<label class="dmc-label">Descrição *</label>')
            _iids['descricao'] = _inp_id('desc')
            val = (_st.get('descricao') or '').replace('"','&quot;')
            ui.html(f'<input id="{_iids["descricao"]}" class="dmc-input" placeholder="ex: Abastecimento" value="{val}">')

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    ui.html('<label class="dmc-label">Valor (R$) *</label>')
                    _iids['valor'] = _inp_id('val')
                    try:
                        _vf = float(_st.get('valor') or 0)
                        _s  = f'{_vf:,.2f}'.replace(',','X').replace('.', ',').replace('X','.')
                        vval = f'R$ {_s}'
                    except (ValueError, TypeError):
                        vval = 'R$ 0,00'
                    ui.html(f'<input id="{_iids["valor"]}" class="dmc-input" type="text" inputmode="numeric" value="{vval}">')
                with ui.element('div'):
                    ui.html('<label class="dmc-label">Vencimento</label>')
                    _iids['data_venc'] = _inp_id('venc')
                    vvenc = (_st.get('data_venc') or '')
                    ui.html(f'<input id="{_iids["data_venc"]}" class="dmc-input" type="date" value="{vvenc}">')

            ui.html('<label class="dmc-label">Categoria</label>')
            _iids['categoria_id'] = _inp_id('cat')
            _new_cat_box_id  = _inp_id('newcatbox')
            _new_cat_nome_id = _inp_id('newcatnome')
            _new_cat_cor_id  = _inp_id('newcatcor')
            cat_cur = _st.get('categoria_id','')
            cat_opts_html = ''.join(
                f'<option value="{k}"{"  selected" if k==cat_cur else ""}>{v}</option>'
                for k, v in cat_opts
            ) + '<option value="__new__">+ Nova categoria…</option>'
            ui.html(
                f'<select id="{_iids["categoria_id"]}" class="dmc-input" style="cursor:pointer">'
                f'{cat_opts_html}</select>'
            )
            # caixa inline para criar categoria nova (oculta por padrão)
            _cor_dots_html = ''.join(
                f'<button type="button" data-cor="{c}" title="{c}" '
                f'style="width:20px;height:20px;border-radius:50%;background:{c};'
                f'border:2px solid transparent;cursor:pointer;outline:none;flex-shrink:0;'
                f'transition:box-shadow .15s,border-color .15s;padding:0"></button>'
                for c in _CAT_COLORS
            )
            ui.html(
                f'<div id="{_new_cat_box_id}" style="display:none;align-items:center;gap:8px;'
                f'margin-top:6px;padding:10px 12px;background:rgba(251,191,36,.05);'
                f'border:1px solid rgba(251,191,36,.2);border-radius:10px">'
                f'<span class="material-icons" style="font-size:15px;color:#FBBF24;flex-shrink:0">label</span>'
                f'<input id="{_new_cat_nome_id}" class="dmc-input" placeholder="Nome da categoria" '
                f'style="flex:1;margin-bottom:0">'
                f'<input type="hidden" id="{_new_cat_cor_id}" value="{_CAT_COLORS[0]}">'
                f'<div style="display:flex;gap:5px;align-items:center;flex-wrap:wrap;flex-shrink:0">'
                f'{_cor_dots_html}'
                f'</div>'
                f'</div>'
            )

            # listener via timer — confiável com html inserido via innerHTML
            async def _setup_cat_toggle():
                await ui.run_javascript(
                    f'(function(){{'
                    f'var sel=document.getElementById("{_iids["categoria_id"]}"),'
                    f'box=document.getElementById("{_new_cat_box_id}");'
                    f'if(!sel||!box)return;'
                    f'sel.addEventListener("change",function(){{'
                    f'box.style.display=this.value==="__new__"?"flex":"none";'
                    f'if(this.value==="__new__"){{'
                    f'var n=document.getElementById("{_new_cat_nome_id}");'
                    f'if(n)n.focus();'
                    f'}}'
                    f'}});'
                    f'var vi=document.getElementById("{_iids["valor"]}");'
                    f'if(vi)vi.oninput=function(){{this.value=maskBRL(this.value);}};'
                    f'var hidCor=document.getElementById("{_new_cat_cor_id}");'
                    f'var dots=box.querySelectorAll("[data-cor]");'
                    f'function selDot(d){{'
                    f'dots.forEach(function(x){{x.style.borderColor="transparent";x.style.boxShadow="none";}});'
                    f'd.style.borderColor="#fff";'
                    f'd.style.boxShadow="0 0 0 3px "+d.getAttribute("data-cor")+"66";'
                    f'if(hidCor)hidCor.value=d.getAttribute("data-cor");'
                    f'}}'
                    f'if(dots.length)selDot(dots[0]);'
                    f'dots.forEach(function(d){{d.addEventListener("click",function(){{selDot(d);}});}});'
                    f'}})()'
                )
            ui.timer(0.15, _setup_cat_toggle, once=True)

            ui.html('<label class="dmc-label">Obra vinculada</label>')
            _iids['obra_id'] = _inp_id('obra')
            obra_cur = _st.get('obra_id','')
            obra_opts_html = ''.join(
                f'<option value="{k}"{"  selected" if k==obra_cur else ""}>{v}</option>'
                for k, v in obra_opts
            )
            ui.html(f'<select id="{_iids["obra_id"]}" class="dmc-input" style="cursor:pointer">{obra_opts_html}</select>')

            if editando:
                ui.html('<label class="dmc-label">Status</label>')
                _iids['status'] = _inp_id('status')
                st_cur = _st.get('status','pendente')
                st_opts_html = ''.join(
                    f'<option value="{k}"{"  selected" if k==st_cur else ""}>{v}</option>'
                    for k, v in [('pendente','Pendente'),('pago','Pago'),('cancelado','Cancelado')]
                )
                ui.html(f'<select id="{_iids["status"]}" class="dmc-input" style="cursor:pointer">{st_opts_html}</select>')

                ui.html('<label class="dmc-label">Data de pagamento</label>')
                _iids['data_pag'] = _inp_id('pag')
                vpag = _st.get('data_pag','')
                ui.html(f'<input id="{_iids["data_pag"]}" class="dmc-input" type="date" value="{vpag}">')

            ui.html('<label class="dmc-label">Observação</label>')
            _iids['observacao'] = _inp_id('obs')
            vobs = (_st.get('observacao') or '').replace('"','&quot;')
            ui.html(f'<input id="{_iids["observacao"]}" class="dmc-input" placeholder="Opcional" value="{vobs}">')

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            async def _salvar():
                js_reads = '{' + ','.join(
                    f'"{k}": (document.getElementById("{v}")||{{}}).value||""'
                    for k, v in _iids.items()
                ) + '}'
                vals = await ui.run_javascript(f'({js_reads})')
                if not isinstance(vals, dict):
                    return
                desc = (vals.get('descricao') or '').strip()
                if not desc:
                    ui.notify('Informe a descrição', color='negative')
                    return
                try:
                    _rv = str(vals.get('valor') or 'R$ 0,00')
                    _rv = _rv.replace('R$','').replace(' ','').replace('.','').replace(',','.')
                    valor = float(_rv or '0')
                except ValueError:
                    valor = 0.0

                obra_id   = vals.get('obra_id','')
                obra_nome = ''
                if obra_id:
                    ob = next((o for o in obras if o['id'] == obra_id), None)
                    if ob:
                        obra_nome = ob.get('cliente_nome','') + (' · ' + ob.get('obra_log','') if ob.get('obra_log') else '')

                cat_id = vals.get('categoria_id','')
                # Criar categoria nova inline se __new__ foi selecionado
                if cat_id == '__new__':
                    new_nome = (await ui.run_javascript(
                        f'(document.getElementById("{_new_cat_nome_id}")||{{}}).value||""'
                    ) or '').strip()
                    new_cor = (await ui.run_javascript(
                        f'(document.getElementById("{_new_cat_cor_id}")||{{}}).value||""'
                    ) or _CAT_COLORS[0])
                    if new_nome:
                        nova_cat = add_categoria_pagar(new_nome, new_cor)
                        cat_id = nova_cat['id']
                    else:
                        cat_id = ''

                if editando:
                    update_conta_pagar(
                        _st['id'],
                        descricao=desc, categoria_id=cat_id, obra_id=obra_id,
                        obra_nome=obra_nome, valor=valor,
                        data_venc=vals.get('data_venc',''),
                        data_pag=vals.get('data_pag',''),
                        status=vals.get('status','pendente'),
                        observacao=vals.get('observacao',''),
                    )
                    ui.notify('Conta atualizada', color='positive')
                else:
                    add_conta_pagar({
                        'descricao': desc, 'categoria_id': cat_id,
                        'obra_id': obra_id, 'obra_nome': obra_nome,
                        'valor': valor, 'data_venc': vals.get('data_venc',''),
                        'observacao': vals.get('observacao',''),
                    })
                    ui.notify('Conta adicionada', color='positive')

                dlg.close()
                if on_save:
                    on_save()

            ui.button('Salvar', icon='save', on_click=_salvar).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary').style(
                'background:rgba(248,113,113,.12);border-color:rgba(248,113,113,.35);color:#F87171'
            )

    dlg.open()


# ═══════════════════════════════════════════════════════════════════════════
# CONTAS A RECEBER
# ═══════════════════════════════════════════════════════════════════════════

def nova_conta_receber_dialog(on_save=None) -> None:
    """Cadastro de conta a receber com parcelas."""
    from services.financeiro import add_conta_receber
    from services.clientes import load_clientes
    from services.obras import load_obras

    clientes = load_clientes()
    obras    = load_obras()
    cli_opts  = [('','— Nenhum —')] + [(c['cpf'], c['nome']) for c in clientes]
    obra_opts = [('','— Nenhuma —')] + [(o['id'], o.get('cliente_nome','') + ' · ' + o.get('obra_log','')) for o in obras]

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(560px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(74,222,128,.08);border:1px solid rgba(74,222,128,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#4ADE80">arrow_circle_up</span></div>'
            )
            with ui.element('div'):
                ui.html('<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text)">Nova Conta a Receber</div>')
                ui.html('<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">Registre valores a receber — parcele se necessário</div>')

        _iids = {}
        def _iid(k): return f'cr-{k}-{id(dlg)}'

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0;display:flex;flex-direction:column;gap:12px'):

            ui.html('<label class="dmc-label">Descrição *</label>')
            _iids['descricao'] = _iid('desc')
            ui.html(f'<input id="{_iids["descricao"]}" class="dmc-input" placeholder="ex: Levantamento topográfico — Lote 12">')

            with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                with ui.element('div'):
                    ui.html('<label class="dmc-label">Valor total (R$) *</label>')
                    _iids['valor_total'] = _iid('vt')
                    ui.html(f'<input id="{_iids["valor_total"]}" class="dmc-input" type="text" inputmode="numeric" value="R$ 0,00">')
                with ui.element('div'):
                    ui.html('<label class="dmc-label">Nº de parcelas</label>')
                    _iids['parcelas'] = _iid('parc')
                    ui.html(f'<input id="{_iids["parcelas"]}" class="dmc-input" type="number" min="1" value="1">')

            ui.html('<label class="dmc-label">Cliente</label>')
            _iids['cliente_cpf'] = _iid('cli')
            cli_opts_html = ''.join(
                f'<option value="{k}">{v}</option>'
                for k, v in cli_opts
            )
            ui.html(f'<select id="{_iids["cliente_cpf"]}" class="dmc-input" style="cursor:pointer">{cli_opts_html}</select>')

            ui.html('<label class="dmc-label">Obra vinculada</label>')
            _iids['obra_id'] = _iid('obra')
            obra_opts_html = ''.join(
                f'<option value="{k}">{v}</option>'
                for k, v in obra_opts
            )
            ui.html(f'<select id="{_iids["obra_id"]}" class="dmc-input" style="cursor:pointer">{obra_opts_html}</select>')

            ui.html('<label class="dmc-label">Vencimento</label>')
            _iids['data_venc'] = _iid('venc')
            ui.html(f'<input id="{_iids["data_venc"]}" class="dmc-input" type="date">')

            ui.html('<label class="dmc-label">Observação</label>')
            _iids['observacao'] = _iid('obs')
            ui.html(f'<input id="{_iids["observacao"]}" class="dmc-input" placeholder="Opcional">')

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;gap:10px;flex-shrink:0'
        ):
            ui.button('Cancelar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

            async def _salvar():
                js_reads = '{' + ','.join(
                    f'"{k}": (document.getElementById("{v}")||{{}}).value||""'
                    for k, v in _iids.items()
                ) + '}'
                vals = await ui.run_javascript(f'({js_reads})')
                if not isinstance(vals, dict):
                    return
                desc = (vals.get('descricao') or '').strip()
                if not desc:
                    ui.notify('Informe a descrição', color='negative')
                    return
                try:
                    _rvt = str(vals.get('valor_total') or 'R$ 0,00')
                    _rvt = _rvt.replace('R$','').replace(' ','').replace('.','').replace(',','.')
                    vt = float(_rvt or '0')
                except ValueError:
                    vt = 0.0
                try:
                    parc = max(1, int(vals.get('parcelas','1') or 1))
                except ValueError:
                    parc = 1

                cpf   = vals.get('cliente_cpf','')
                cli   = next((c for c in clientes if c['cpf'] == cpf), None)
                cnome = cli['nome'] if cli else ''

                obra_id   = vals.get('obra_id','')
                obra_nome = ''
                if obra_id:
                    ob = next((o for o in obras if o['id'] == obra_id), None)
                    if ob:
                        obra_nome = ob.get('cliente_nome','') + (' · ' + ob.get('obra_log','') if ob.get('obra_log') else '')

                add_conta_receber({
                    'descricao': desc, 'cliente_nome': cnome, 'cliente_cpf': cpf,
                    'obra_id': obra_id, 'obra_nome': obra_nome,
                    'valor_total': vt, 'parcelas': parc,
                    'data_venc': vals.get('data_venc',''),
                })
                ui.notify('Conta adicionada', color='positive')
                dlg.close()
                if on_save:
                    on_save()

            ui.button('Salvar', icon='save', on_click=_salvar).props(
                'unelevated no-caps'
            ).classes('dmc-btn dmc-btn-primary')

    _vt_id = _iids['valor_total']
    async def _setup_receber_masks():
        await ui.run_javascript(
            f'(function(){{'
            f'var vi=document.getElementById("{_vt_id}");'
            f'if(vi)vi.oninput=function(){{this.value=maskBRL(this.value);}};'
            f'}})()'
        )
    ui.timer(0.15, _setup_receber_masks, once=True)

    dlg.open()


def registrar_pagamento_dialog(conta: dict, on_save=None, modo: str = "pagar") -> None:
    """Registra parcelas / exibe histórico de pagamentos de uma conta a receber."""
    from services.financeiro import add_pagamento, delete_pagamento

    vt      = float(conta.get('valor_total', 0))
    vp      = float(conta.get('valor_pago', 0))
    restante = round(vt - vp, 2)
    parc    = int(conta.get('parcelas', 1))
    pags    = conta.get('pagamentos', [])

    def _fmt(v): return f'R$ {float(v):,.2f}'.replace(',','X').replace('.', ',').replace('X','.')

    with ui.dialog().props('persistent') as dlg, ui.card().style(
        'background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;'
        'border-radius:18px!important;padding:0;'
        'width:min(580px,97vw)!important;max-height:94vh;'
        'display:flex;flex-direction:column;color:var(--dmc-text)!important;position:relative;'
    ):
        ui.button(icon='close', on_click=dlg.close).props('flat round dense').style(
            'color:var(--dmc-muted);position:absolute;top:12px;right:12px;z-index:10;'
        )

        # Header
        with ui.element('div').style(
            'padding:18px 24px;border-bottom:1px solid var(--dmc-b1);'
            'display:flex;align-items:center;gap:14px;flex-shrink:0;padding-right:52px;'
        ):
            ui.html(
                '<div style="width:40px;height:40px;border-radius:10px;flex-shrink:0;'
                'background:rgba(74,222,128,.08);border:1px solid rgba(74,222,128,.25);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#4ADE80">payments</span></div>'
            )
            with ui.element('div').style('flex:1;min-width:0'):
                ui.html(f'<div style="font:700 16px var(--dmc-fd);color:var(--dmc-text);overflow:hidden;text-overflow:ellipsis;white-space:nowrap">{conta.get("descricao","—")}</div>')
                ui.html(f'<div style="font:12px var(--dmc-fm);color:var(--dmc-muted2);margin-top:1px">{conta.get("cliente_nome","—")}</div>')

        with ui.element('div').style('padding:20px 24px;overflow-y:auto;flex:1;min-height:0;display:flex;flex-direction:column;gap:14px'):

            # Resumo financeiro
            pct = int(min(100, (vp/vt*100))) if vt else 0
            st  = conta.get('status','aberto')
            from pages.financeiro import _STATUS_RECEBER
            cor, bg = _STATUS_RECEBER.get(st, ('#8BAA8B','rgba(139,170,139,.08)'))

            ui.html(
                f'<div style="background:rgba(255,255,255,.03);border:1px solid var(--dmc-b1);'
                f'border-radius:12px;padding:14px 16px;display:flex;flex-direction:column;gap:8px">'
                f'<div style="display:flex;justify-content:space-between;align-items:center">'
                f'<span style="font:11px var(--dmc-mono);color:var(--dmc-muted2)">TOTAL</span>'
                f'<span style="font:700 16px var(--dmc-fd);color:#DCE8DC">{_fmt(vt)}</span>'
                f'</div>'
                f'<div style="height:6px;border-radius:3px;background:rgba(255,255,255,.06);overflow:hidden">'
                f'<div style="width:{pct}%;height:100%;background:{cor};border-radius:3px;transition:width .3s"></div>'
                f'</div>'
                f'<div style="display:flex;justify-content:space-between">'
                f'<span style="font:11px var(--dmc-mono);color:{cor}">Recebido: {_fmt(vp)}</span>'
                f'<span style="font:11px var(--dmc-mono);color:var(--dmc-muted2)">Restante: {_fmt(restante)}</span>'
                f'</div>'
                f'<div style="display:flex;align-items:center;gap:6px">'
                f'<span style="font:700 9px var(--dmc-mono);background:{bg};color:{cor};'
                f'padding:2px 8px;border-radius:4px;border:1px solid {cor}44">{st.upper()}</span>'
                f'<span style="font:11px var(--dmc-mono);color:var(--dmc-muted2)">'
                f'{len(pags)}/{parc} parcela{"s" if parc!=1 else ""} registrada{"s" if len(pags)!=1 else ""}'
                f'</span></div>'
                f'</div>'
            )

            # Histórico de pagamentos
            pags_area = ui.element('div').style('display:flex;flex-direction:column;gap:6px')

            def _draw_pags():
                pags_area.clear()
                with pags_area:
                    if not pags:
                        ui.html('<div style="text-align:center;padding:20px;font:11px var(--dmc-mono);color:var(--dmc-muted2)">Nenhum pagamento registrado</div>')
                        return
                    for p in pags:
                        pid = p['id']
                        with ui.element('div').style(
                            'display:flex;align-items:center;gap:10px;'
                            'padding:10px 14px;background:rgba(255,255,255,.03);'
                            'border:1px solid var(--dmc-b1);border-radius:10px;'
                        ):
                            ui.html('<span class="material-icons" style="font-size:16px;color:#4ADE80;flex-shrink:0">check_circle</span>')
                            with ui.element('div').style('flex:1;min-width:0'):
                                ui.html(f'<div style="font:600 13px var(--dmc-mono);color:#4ADE80">{_fmt(p["valor"])}</div>')
                                if p.get('observacao'):
                                    ui.html(f'<div style="font:11px var(--dmc-fm);color:var(--dmc-muted2)">{p["observacao"]}</div>')
                            ui.html(f'<span style="font:11px var(--dmc-mono);color:var(--dmc-muted2);flex-shrink:0">{p.get("data","—")}</span>')
                            del_btn = ui.element('button').style(
                                'width:26px;height:26px;border-radius:5px;cursor:pointer;'
                                'background:transparent;border:1px solid rgba(248,113,113,.2);'
                                'display:inline-flex;align-items:center;justify-content:center;color:#F87171'
                            )
                            with del_btn:
                                ui.html('<span class="material-icons" style="font-size:13px">delete</span>')
                            def _del(pid_=pid):
                                delete_pagamento(pid_)
                                updated = [p for p in pags if p['id'] != pid_]
                                pags.clear()
                                pags.extend(updated)
                                _draw_pags()
                                if on_save:
                                    on_save()
                            del_btn.on('click', _del)

            _draw_pags()

            # Formulário de novo pagamento (só em modo pagar e se não quitado)
            if modo == 'pagar' and st != 'quitado':
                ui.html(
                    '<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);'
                    'letter-spacing:.12em;text-transform:uppercase;border-top:1px solid var(--dmc-b1);'
                    'padding-top:14px">Novo Pagamento</div>'
                )
                _pv_id  = f'pv-val-{id(dlg)}'
                _pd_id  = f'pv-dat-{id(dlg)}'
                _po_id  = f'pv-obs-{id(dlg)}'
                sugestao = round(restante / max(1, parc - len(pags)), 2) if restante > 0 else restante

                with ui.element('div').style('display:grid;grid-template-columns:1fr 1fr;gap:14px'):
                    with ui.element('div'):
                        ui.html('<label class="dmc-label">Valor (R$) *</label>')
                        ui.html(f'<input id="{_pv_id}" class="dmc-input" type="number" step="0.01" min="0.01" max="{round(restante,2)}" value="{sugestao if sugestao > 0 else ""}" placeholder="0,00">')
                    with ui.element('div'):
                        ui.html('<label class="dmc-label">Data</label>')
                        ui.html(f'<input id="{_pd_id}" class="dmc-input" type="date">')
                ui.html('<label class="dmc-label">Observação</label>')
                ui.html(f'<input id="{_po_id}" class="dmc-input" placeholder="Opcional">')

                async def _reg_pag():
                    vals = await ui.run_javascript(
                        f'({{ val:(document.getElementById("{_pv_id}")||{{}}).value||"",'
                        f'dat:(document.getElementById("{_pd_id}")||{{}}).value||"",'
                        f'obs:(document.getElementById("{_po_id}")||{{}}).value||""}})'
                    )
                    try:
                        v = float(str(vals.get('val','0')).replace(',','.'))
                    except Exception:
                        v = 0.0
                    if v <= 0:
                        ui.notify('Informe um valor válido', color='negative')
                        return
                    novo_pag = add_pagamento(conta['id'], v, vals.get('dat',''), vals.get('obs',''))
                    pags.append(novo_pag)
                    # Atualiza estado local
                    nonlocal vp, restante, st
                    vp       = round(vp + v, 2)
                    restante = round(vt - vp, 2)
                    st       = 'quitado' if vp >= vt else 'parcial'
                    _draw_pags()
                    ui.notify('Pagamento registrado', color='positive')
                    if on_save:
                        on_save()
                    if restante <= 0:
                        dlg.close()

                with ui.element('div').style('display:flex;justify-content:flex-end'):
                    ui.button('Registrar Pagamento', icon='payments', on_click=_reg_pag).props(
                        'unelevated no-caps'
                    ).classes('dmc-btn dmc-btn-primary')

        with ui.element('div').style(
            'padding:14px 24px;border-top:1px solid var(--dmc-b1);'
            'display:flex;justify-content:flex-end;flex-shrink:0'
        ):
            ui.button('Fechar', on_click=dlg.close).props('flat no-caps').classes('dmc-btn dmc-btn-ghost')

    dlg.open()


# ── Lixeira Financeira ────────────────────────────────────────────────────────

def lixeira_financeiro_dialog(on_restore=None) -> None:
    from services.financeiro import (
        load_contas_pagar_deletadas, load_contas_receber_deletadas,
        restore_conta_pagar, restore_conta_receber,
        purge_conta_pagar, purge_conta_receber,
    )

    with ui.dialog() as dlg, ui.card().style(
        "background:var(--dmc-bg2)!important;border:1px solid var(--dmc-b2)!important;"
        "border-radius:18px!important;padding:0;"
        "width:min(860px,97vw)!important;max-height:88vh;"
        "display:flex;flex-direction:column;color:var(--dmc-text)!important;"
    ):
        with ui.element("div").style(
            "padding:18px 24px;border-bottom:1px solid var(--dmc-b1);"
            "display:flex;align-items:center;gap:14px;flex-shrink:0"
        ):
            ui.html(
                '<div style="width:38px;height:38px;border-radius:10px;flex-shrink:0;'
                'background:rgba(248,113,113,.08);border:1px solid rgba(248,113,113,.3);'
                'display:flex;align-items:center;justify-content:center;">'
                '<span class="material-icons" style="font-size:20px;color:#F87171">delete_outline</span></div>'
            )
            with ui.element("div").style("flex:1"):
                ui.html('<div style="font:700 15px var(--dmc-fd);color:var(--dmc-text)">Lixeira Financeira</div>')
                ui.html(
                    '<div style="font:11px var(--dmc-fm);color:var(--dmc-muted2);margin-top:2px">'
                    'Contas excluídas — restaure ou exclua permanentemente</div>'
                )
            ui.button(icon="close", on_click=dlg.close).props("flat round dense").style(
                "color:var(--dmc-muted)"
            )

        content_area = ui.element("div").style("padding:18px 24px;overflow-y:auto;flex:1")

        def _brl(v):
            return f"R$ {float(v or 0):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

        def _render_lixeira():
            content_area.clear()
            pagar   = load_contas_pagar_deletadas()
            receber = load_contas_receber_deletadas()
            with content_area:
                if not pagar and not receber:
                    ui.html(
                        '<div style="text-align:center;padding:60px 0;color:var(--dmc-muted2)">'
                        '<span class="material-icons" style="font-size:52px;opacity:.2;'
                        'display:block;margin-bottom:12px">delete_outline</span>'
                        '<div style="font:13px var(--dmc-fm)">Nenhuma conta na lixeira</div>'
                        '</div>'
                    )
                    return

                for titulo, itens, is_pagar in [
                    ("Contas a Pagar", pagar, True),
                    ("Contas a Receber", receber, False),
                ]:
                    if not itens:
                        continue
                    ui.html(
                        f'<div style="font:600 11px var(--dmc-mono);color:var(--dmc-muted2);'
                        f'letter-spacing:.14em;text-transform:uppercase;margin-bottom:8px">{titulo}</div>'
                    )
                    for c in itens:
                        cid  = c["id"]
                        desc = c.get("descricao", "—")
                        val  = _brl(c.get("valor") or c.get("valor_total", 0))
                        venc = c.get("data_venc", "—")
                        del_em = c.get("deletado_em", "—")
                        with ui.element("div").style(
                            "background:var(--dmc-bg3);border:1px solid var(--dmc-b1);border-radius:10px;"
                            "padding:10px 14px;margin-bottom:6px;"
                            "display:flex;align-items:center;gap:12px"
                        ):
                            with ui.element("div").style("flex:1;min-width:0"):
                                ui.html(
                                    f'<div style="font:500 13px var(--dmc-fm);color:var(--dmc-text);'
                                    f'overflow:hidden;text-overflow:ellipsis;white-space:nowrap">{desc}</div>'
                                    f'<div style="font:11px var(--dmc-mono);color:var(--dmc-muted2);margin-top:3px">'
                                    f'{val} · Venc.: {venc} · Excluída: {del_em}</div>'
                                )
                            with ui.element("div").style("display:flex;gap:6px;flex-shrink:0"):
                                rst = ui.element("button").classes("dmc-btn dmc-btn-secondary dmc-btn-sm")
                                with rst:
                                    ui.html(
                                        '<span class="material-icons" style="font-size:13px">settings_backup_restore</span>'
                                        "<span>Restaurar</span>"
                                    )

                                def _restore(cid=cid, ip=is_pagar):
                                    if ip:
                                        restore_conta_pagar(cid)
                                    else:
                                        restore_conta_receber(cid)
                                    ui.notify("Conta restaurada.", type="positive")
                                    if on_restore:
                                        on_restore()
                                    _render_lixeira()

                                rst.on("click", _restore)

                                prg = ui.element("button").classes("dmc-btn dmc-btn-danger dmc-btn-sm")
                                with prg:
                                    ui.html(
                                        '<span class="material-icons" style="font-size:13px">delete_forever</span>'
                                        "<span>Excluir</span>"
                                    )

                                def _purge(cid=cid, ip=is_pagar):
                                    if ip:
                                        purge_conta_pagar(cid)
                                    else:
                                        purge_conta_receber(cid)
                                    ui.notify("Excluído permanentemente.", type="warning")
                                    _render_lixeira()

                                prg.on("click", _purge)

                    ui.element("div").style("height:10px")

        _render_lixeira()

        with ui.element("div").style(
            "padding:12px 24px;border-top:1px solid var(--dmc-b1);"
            "display:flex;justify-content:flex-end;flex-shrink:0"
        ):
            ui.button("Fechar", on_click=dlg.close).props("flat no-caps").classes("dmc-btn dmc-btn-ghost")

    dlg.open()
